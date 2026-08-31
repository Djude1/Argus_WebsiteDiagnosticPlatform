"""Word（.docx）健檢報告產生器。

報告的讀者是網站主，不是資安工程師（見 docs/scan-report-quality-audit-2026-08-30.md）。
所以這裡的原則是：

- 內部識別碼（rule_id、evidence_source、evidence_type）不進正文，rule_id 收進附錄
  的技術索引供工程師查用。
- 每一筆發現固定回答四個問題：問題是什麼 / 為什麼要在意 / 怎麼修 / 修好了怎麼確認。
  舊版依 severity 給同一個結構三種標題（風險描述／改善重點／建議優化），讀者會以為
  是三種不同的東西。
- 名詞解釋只列這份報告裡真的出現過的術語，不是貼一份固定清單。
- 資訊要有結構（表格、分頁、頁碼、嚴重度顏色），不是 300 段純文字流。

內容邊界（哪些欄位不得寫進報告）見 backend/apps/scans/CLAUDE.md「報告內容契約」。
"""

import hmac
from collections import OrderedDict
from hashlib import sha256
from pathlib import Path

from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from apps.scans.models import Finding, ReportVerification, ScanJob
from apps.scans.scan_plan import build_scan_execution_plan
from apps.scans.security.redaction import redact_pii_in_text

# --- 品牌與嚴重度配色 -------------------------------------------------
# frontend/src/styles.css 的 token 是為深色背景設計的，--argus-cyan (#38bdf8)
# 印在白紙上對比不足，所以標題改用 --argus-cyan-deep，cyan 只當強調線。
ARGUS_NAVY = "0A1535"        # --argus-navy-800
ARGUS_CYAN = "38BDF8"        # --argus-cyan（強調線用）
ARGUS_CYAN_DEEP = "0C4A6E"   # --argus-cyan-deep（白底可讀的標題色）
ARGUS_MUTED = "5B6B7C"

SEVERITY_COLOR = {
    "critical": "B02418",
    "high": "C4600F",
    "medium": "8A6A00",
    "low": "1F6591",
    "info": "5B4B8A",
}

SEVERITY_DISPLAY = {
    "critical": "嚴重風險",
    "high": "高風險",
    "medium": "中風險",
    "low": "低風險",
    "info": "資訊提示",
}

CATEGORY_DISPLAY = {
    "seo": "SEO 搜尋引擎最佳化",
    "aeo": "AEO 問答引擎最佳化",
    "geo": "GEO 生成式引擎最佳化",
    "security": "資訊安全",
    "ux": "使用者體驗",
}

SCORE_BANDS = [
    (80, "良好", "持續維持即可，建議定期複檢。"),
    (60, "需改善", "有幾項體質問題值得排入維護排程。"),
    (40, "建議儘快處理", "累積的問題已可能影響流量或安全，建議近期處理。"),
    (0, "需優先處理", "存在較高風險的項目，建議優先安排修補。"),
]

# 只列這份報告裡真的出現過的術語。key 會拿去比對報告文字。
GLOSSARY = {
    "HSTS": "強制瀏覽器之後一律以加密連線（HTTPS）連到你的網站，避免被降級成未加密連線。",
    "CSP": "內容安全政策。告訴瀏覽器這個網頁只能載入哪些來源的程式與資源，"
           "用來擋掉被植入的惡意腳本。",
    "CSRF": "跨站請求偽造。攻擊者誘導已登入的使用者在不知情下送出操作（例如改密碼、轉帳）。",
    "SPF": "在 DNS 上公告「哪些伺服器有資格用我的網域寄信」，別人冒名寄信時較容易被擋下。",
    "DMARC": "搭配 SPF 使用，告訴收信方「查到冒名信件時要怎麼處理」（放行、隔離或退回）。",
    "DNSSEC": "為 DNS 查詢結果加上數位簽章，避免有人竄改網域解析把訪客導到假網站。",
    "SRI": "子資源完整性。為外部載入的 JS/CSS 加上指紋，檔案被竄改時瀏覽器會拒絕執行。",
    "CORS": "跨來源資源共用。控制哪些其他網站可以用瀏覽器讀取你的 API 回應。",
    "X-Frame-Options": "限制你的網頁能不能被別的網站嵌入框架，用來防止點擊劫持。",
    "X-Content-Type-Options": "要求瀏覽器嚴格照宣告的檔案類型處理，不要自行猜測。",
    "canonical": "標準網址。同一份內容有多個網址時，指定哪一個才是正式版本，避免搜尋權重被分散。",
    "JSON-LD": "一種結構化資料格式，讓搜尋引擎與 AI 更準確理解頁面在講什麼。",
    "robots.txt": "放在網站根目錄的檔案，用來告訴各種爬蟲哪些路徑可以抓、哪些不要抓。",
    "llms.txt": "類似 robots.txt 的新興慣例，用來對 AI 模型說明網站內容與可引用範圍。",
    "WAF": "網站應用防火牆。擋在網站前面過濾惡意請求的防護層。",
    "PII": "個人可識別資訊，例如 Email、手機號碼、身分證字號。",
    "TLS": "網路傳輸加密協定，HTTPS 底層用的就是它（早期稱為 SSL）。",
    "OWASP": "國際公認的網站安全風險分類，A01~A10 代表十大類風險。",
    "CWE": "國際通用的軟體弱點編號系統，用來精確指出是哪一種弱點。",
}

# 「為什麼要在意」：只陳述該分類的一般性後果，不臆測個案細節。
CATEGORY_IMPACT = {
    "security": "這類問題會被攻擊者利用，可能導致網站被入侵、使用者資料外洩，"
                "或你的網域被冒用來寄送釣魚信件，連帶損害品牌信任。",
    "seo": "這類問題會讓搜尋引擎較難正確理解與收錄你的頁面，"
           "潛在客戶用關鍵字搜尋時，你的網站可能排在競爭對手後面。",
    "aeo": "這類問題會讓 AI 助理在回答使用者提問時，較難引用你的內容，"
           "等於在新的搜尋入口上失去曝光機會。",
    "geo": "這類問題會讓生成式搜尋引擎難以擷取與理解你的頁面主題，"
           "影響你的內容被 AI 摘要與推薦的機會。",
    "ux": "這類問題會讓訪客在瀏覽或操作時遇到阻礙，"
          "直接反映在跳出率與轉換率上。",
}

# info 有兩種：正向指標（「探針被 WAF 擋下，代表防護有效」）與可改的小問題
# （缺 X-Content-Type-Options、缺 canonical）。scan 28 的 5 個 info 裡只有 1 個
# 是正向。所以這段文字兩者都要成立——既不能套 CATEGORY_IMPACT 那套「會被攻擊者
# 利用」（對正向指標完全相反），也不能宣告「不需要任何修補動作」（對另外 4 個
# 是叫人別管一個其實可以改的東西，而且下一行就印出修補方式，自相矛盾）。
INFO_NOTE = (
    "這是一項影響較小的觀察項目，不屬於需要立即處理的風險。"
    "若下方列有修補方式，可視情況安排。"
)

SEVERITY_URGENCY = {
    "critical": "這是本次掃描中最高等級的風險，建議立即處理。",
    "high": "建議優先排入處理，不要拖過本次維護週期。",
    "medium": "建議排入近期的維護排程。",
    "low": "屬於體質項目，可與其他改善一起處理。",
}

CATEGORY_VERIFY = {
    "security": "修補後重新執行一次 Argus 掃描，確認此項目不再出現。"
                "若要立即自行確認，可請你的網站維護人員依上方「怎麼修」的步驟逐項檢查。",
    "seo": "修補後重新執行一次 Argus 掃描確認此項目消失，"
           "並可用 Google Search Console 觀察後續的索引狀態。",
    "aeo": "修補後重新執行一次 Argus 掃描確認此項目消失。",
    "geo": "修補後重新執行一次 Argus 掃描確認此項目消失。",
    "ux": "修補後重新執行一次 Argus 掃描確認此項目消失，並請實際操作一次該流程。",
}

DISCLAIMER = (
    "免責聲明：本報告僅反映掃描當下、從網際網路可觀測到的外部特徵，"
    "不等同完整滲透測試或原始碼稽核，也不構成法律或合規意見。"
    "未列出的項目不代表不存在風險。實際修補請由具備權限的維運人員評估後執行。"
)

_CJK_FONT = "Microsoft JhengHei"


def build_report_number(scan_job: ScanJob) -> str:
    """產生對外揭露的報告編號：ARGUS-{掃描編號}-{日期}-{4 碼驗證碼}。

    驗證碼用 SECRET_KEY 做 HMAC，讓編號無法被憑空捏造出「看起來合理」的值；
    只取 4 碼是因為它防的是隨手偽造，真正的比對靠查驗端點與內容雜湊。

    刻意只用 scan_job 的固定欄位（不含產生時間），所以**重新產生報告時編號不變**
    ——報告一旦交付就可能被轉寄存檔，換編號會讓已流出的副本失效。
    """
    issued = scan_job.completed_at or scan_job.created_at or timezone.now()
    token = hmac.new(
        settings.SECRET_KEY.encode("utf-8"),
        f"argus-report:{scan_job.pk}".encode(),
        sha256,
    ).hexdigest()[:4].upper()
    return f"ARGUS-{scan_job.pk}-{issued.strftime('%Y%m%d')}-{token}"


def get_severity_display(severity: str) -> str:
    return SEVERITY_DISPLAY.get(severity, severity or "未知")


def mask_pii_evidence(text: str) -> str:
    """報告展示層遮罩：evidence 含原始個資（email/手機/身分證/信用卡）就地遮罩。

    這份 .docx 會被下載、轉寄、存檔，直接印出原始內容有明確合規風險。只保留頭尾供
    人工比對，不修改 DB 內的原始 Finding 記錄。

    對「所有」finding 的 evidence 都套用（不靠 rule_id 前綴判斷是否為 PII finding）：
    除了 scanners.py::analyze_data_exposure() 產生的 SECURITY_PII_* finding，
    security/exposure_scanner.py 的敏感檔案外洩 finding 也會把命中檔案的原始內容
    片段放進 evidence，一樣可能含未遮罩個資，用 rule_id 白名單很容易漏掉這類來源；
    正則對不含 PII 樣式的文字（如 header 名稱、URL）是無操作，不會誤傷正常內容。
    """
    return redact_pii_in_text(text)


def _group_findings_for_report(findings) -> list[dict]:
    """同一個 rule_id 的 finding 合併成一筆，受影響頁面收斂成清單。

    合併鍵只用 rule_id。rule_id 由 scanners._default_rule_id() 從 category + title
    的雜湊產生，同一種問題不論出現在哪一頁都一致。舊版把 evidence 也放進鍵裡，
    但 evidence 帶的是該頁專屬內容（例如那一頁實際的 title 文字），於是同一個問題
    出現在 N 頁就被拆成 N 筆顯示——scan 25 的報告因此把「Meta title 長度不理想」
    列了 4 次、「核心內容高度依賴 JavaScript 渲染」列了 3 次。

    rule_id 為空時退回 finding.pk，避免不同問題只因為「都沒有 rule_id」被錯誤
    合併成一筆。只影響 .docx 呈現順序與分組，不改資料庫裡的原始 Finding 記錄。
    """
    groups: OrderedDict[str, dict] = OrderedDict()
    for finding in findings:
        key = finding.rule_id or f"_finding:{finding.pk}"
        group = groups.get(key)
        if group is None:
            group = {"finding": finding, "pages": []}
            groups[key] = group
        page_label = finding.page.final_url if finding.page else "站台層級"
        if page_label not in group["pages"]:
            group["pages"].append(page_label)
    return list(groups.values())


# --- 低階排版工具 -----------------------------------------------------


def _styled_run(paragraph, text: str, *, size=None, bold=False, color=None):
    """加一段文字並設好中文字型。

    python-docx 的 run.font.name 只設拉丁字型，中文字要另外設 w:eastAsia，
    否則 Word 會退回預設字型，中英文混排看起來會不一致。
    """
    run = paragraph.add_run(text)
    run.font.name = _CJK_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), _CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _heading(document, text: str, level: int):
    """章節標題。用 Heading 樣式維持 Word 的大綱結構，但改成品牌色。"""
    heading = document.add_heading(level=level)
    _styled_run(heading, text, bold=True, color=ARGUS_CYAN_DEEP)
    return heading


def _accent_rule(document):
    """一條 cyan 強調線，用來分隔封面區塊。"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(paragraph, "━" * 24, color=ARGUS_CYAN)
    return paragraph


def _kv_table(document, rows: list[tuple[str, str]]):
    """兩欄的「項目 / 內容」表格。"""
    table = document.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        cells = table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], label, bold=True)
        _styled_run(cells[1].paragraphs[0], value)
    return table


def _header_row(table, labels: list[str]):
    for cell, label in zip(table.rows[0].cells, labels, strict=False):
        _styled_run(cell.paragraphs[0], label, bold=True, color=ARGUS_CYAN_DEEP)


def _add_header_footer(document, scan_job: ScanJob, report_number: str) -> None:
    section = document.sections[0]

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(
        header_paragraph,
        f"ARGUS 網站健檢報告　|　{scan_job.origin}",
        size=9, color=ARGUS_MUTED,
    )

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(footer_paragraph, "Argus 授權式網站健檢　|　第 ", size=9, color=ARGUS_MUTED)
    # 頁碼要用 Word 的 field code，寫死數字會在分頁變動時失準
    run = footer_paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    _styled_run(
        footer_paragraph, f" 頁　|　{report_number}", size=9, color=ARGUS_MUTED,
    )


def _score_band(score) -> tuple[str, str]:
    if not isinstance(score, (int, float)):
        return "尚未產生", ""
    for threshold, label, advice in SCORE_BANDS:
        if score >= threshold:
            return label, advice
    return "需優先處理", ""


# --- 章節 -------------------------------------------------------------


def _add_cover(document, scan_job: ScanJob, report_number: str) -> None:
    # 有 PNG 就放圖，沒有就用字標。刻意不引入 SVG 轉檔套件（cairosvg 有系統函式庫
    # 相依，會拖累 CI 與 Docker build），把 frontend/public/argus-logo.png 放進去
    # 就會自動生效。
    logo_path = Path(settings.PROJECT_ROOT) / "frontend" / "public" / "argus-logo.png"
    if logo_path.exists():
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(1.4))
    else:
        for _ in range(3):
            document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(title, "ARGUS", size=44, bold=True, color=ARGUS_NAVY)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(subtitle, "網站健檢報告", size=22, bold=True, color=ARGUS_CYAN_DEEP)

    _accent_rule(document)
    document.add_paragraph()

    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(target, scan_job.normalized_url, size=13, bold=True)

    band, _ = _score_band(scan_job.overall_score)
    score_paragraph = document.add_paragraph()
    score_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_text = (
        f"整體分數 {scan_job.overall_score} / 100　（{band}）"
        if scan_job.overall_score is not None
        else "整體分數：尚未產生"
    )
    _styled_run(score_paragraph, score_text, size=15, bold=True, color=ARGUS_CYAN_DEEP)

    document.add_paragraph()
    completed = (
        scan_job.completed_at.strftime("%Y-%m-%d %H:%M:%S")
        if scan_job.completed_at else "—"
    )
    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(
        meta,
        f"掃描完成：{completed}\n"
        f"報告產生：{timezone.localtime().strftime('%Y-%m-%d %H:%M:%S')}",
        size=10, color=ARGUS_MUTED,
    )

    document.add_paragraph()
    number_paragraph = document.add_paragraph()
    number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(number_paragraph, f"報告編號　{report_number}", size=11, bold=True)
    verify_paragraph = document.add_paragraph()
    verify_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(
        verify_paragraph,
        "可至 Argus 網站的「報告查驗」頁（/verify）輸入上方編號，核對本報告的真偽與內容指紋。",
        size=9, color=ARGUS_MUTED,
    )
    document.add_page_break()


def _add_table_of_contents(document, grouped_findings) -> None:
    _heading(document, "目錄", level=1)
    sections = [
        ("1", "摘要", "整體分數、各分類分數與發現項目統計"),
        ("2", "掃描範圍", "這份報告涵蓋了什麼、沒涵蓋什麼"),
        ("3", "優先改善建議", "最值得先處理的項目"),
        ("4", "發現項目", f"逐項說明，共 {len(grouped_findings)} 項"),
        ("5", "附錄", "授權聲明、名詞解釋、技術索引"),
    ]
    table = document.add_table(rows=len(sections) + 1, cols=3)
    table.style = "Table Grid"
    _header_row(table, ["", "章節", "內容"])
    for index, (number, name, desc) in enumerate(sections, start=1):
        cells = table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], number, bold=True)
        _styled_run(cells[1].paragraphs[0], name, bold=True)
        _styled_run(cells[2].paragraphs[0], desc)
    document.add_page_break()


def _add_summary(document, scan_job: ScanJob, grouped_findings) -> None:
    _heading(document, "1　摘要", level=1)
    document.add_paragraph()
    _styled_run(
        document.add_paragraph(),
        "整體分數是下列「已評估」分類分數的平均，標示「未評估」的分類不納入計算。"
        "同一個問題出現在多個頁面只扣一次分；資訊提示屬正向或中性訊息，不扣分。",
    )

    _heading(document, "各分類分數", level=2)
    category_scores = scan_job.category_scores or {}
    table = document.add_table(rows=len(Finding.Category.values) + 1, cols=2)
    table.style = "Table Grid"
    _header_row(table, ["分類", "分數"])
    for index, category in enumerate(Finding.Category.values, start=1):
        cells = table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], CATEGORY_DISPLAY.get(category, category.upper()))
        score = category_scores.get(category)
        if isinstance(score, (int, float)):
            _styled_run(cells[1].paragraphs[0], f"{score} 分", bold=True)
        else:
            _styled_run(
                cells[1].paragraphs[0], "未評估（本次未執行此項檢查）", color=ARGUS_MUTED
            )

    _add_previous_scan_comparison(document, scan_job, grouped_findings)

    _heading(document, "分數怎麼看", level=2)
    band_table = document.add_table(rows=len(SCORE_BANDS) + 1, cols=3)
    band_table.style = "Table Grid"
    _header_row(band_table, ["分數", "評級", "建議"])
    ranges = ["80 - 100", "60 - 79", "40 - 59", "0 - 39"]
    for index, ((_, label, advice), score_range) in enumerate(
        zip(SCORE_BANDS, ranges, strict=True), start=1
    ):
        cells = band_table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], score_range)
        _styled_run(cells[1].paragraphs[0], label, bold=True)
        _styled_run(cells[2].paragraphs[0], advice)

    _heading(document, "發現項目統計", level=2)
    counts = {severity: 0 for severity in SEVERITY_DISPLAY}
    for item in grouped_findings:
        counts[item["finding"].severity] = counts.get(item["finding"].severity, 0) + 1
    stat_table = document.add_table(rows=len(SEVERITY_DISPLAY) + 1, cols=2)
    stat_table.style = "Table Grid"
    _header_row(stat_table, ["嚴重度", "項目數"])
    for index, severity in enumerate(SEVERITY_DISPLAY, start=1):
        cells = stat_table.rows[index].cells
        _styled_run(
            cells[0].paragraphs[0], get_severity_display(severity),
            bold=True, color=SEVERITY_COLOR[severity],
        )
        _styled_run(cells[1].paragraphs[0], str(counts.get(severity, 0)))
    document.add_page_break()


class _SectionNumber:
    """附錄小節的流水號。章節會依資料有無而出現或消失，號碼不能寫死。"""

    def __init__(self, prefix: str) -> None:
        self._prefix = prefix
        self._index = 0

    def next(self) -> str:
        self._index += 1
        return f"{self._prefix}.{self._index}"


def _previous_completed_scan(scan_job: ScanJob):
    """同一位使用者、同一個 origin 的上一次完成掃描。

    只看本人的掃描：同一個網址可能被不同使用者掃過，拿別人的結果當「前次」
    既不合理也會洩漏他人的掃描存在。未完成或沒有分數的掃描不算。
    """
    if scan_job.completed_at is None:
        return None
    return (
        ScanJob.objects.filter(
            user_id=scan_job.user_id,
            origin=scan_job.origin,
            status=ScanJob.Status.COMPLETED,
            overall_score__isnull=False,
            completed_at__lt=scan_job.completed_at,
        )
        .exclude(pk=scan_job.pk)
        .order_by("-completed_at")
        .first()
    )


def _add_previous_scan_comparison(document, scan_job: ScanJob, grouped_findings) -> None:
    """與前次掃描比較。第一次掃描這個網站時整段不出現。

    「上次 39 分、這次 65 分，修好了 3 項」是回訪使用者最想看到的東西，
    資料早就在 DB 裡（同 origin 的歷史掃描），舊版報告一個字都沒用。
    """
    previous = _previous_completed_scan(scan_job)
    if previous is None:
        return

    _heading(document, "與前次掃描比較", level=2)
    current_score = scan_job.overall_score
    delta_text = "—"
    if isinstance(current_score, int) and isinstance(previous.overall_score, int):
        delta = current_score - previous.overall_score
        delta_text = f"{delta:+d} 分" if delta else "持平"
    _kv_table(document, [
        (
            "前次掃描",
            f"{previous.completed_at.strftime('%Y-%m-%d')}　{previous.overall_score} 分",
        ),
        ("本次掃描", f"{current_score if current_score is not None else '—'} 分"),
        ("變化", delta_text),
    ])

    previous_rules = {
        rule for rule in previous.findings.values_list("rule_id", flat=True) if rule
    }
    current_map = {
        item["finding"].rule_id: item["finding"].title
        for item in grouped_findings
        if item["finding"].rule_id
    }
    previous_titles = {
        rule: title
        for rule, title in previous.findings.values_list("rule_id", "title")
        if rule
    }

    resolved = sorted(previous_rules - set(current_map))
    appeared = sorted(set(current_map) - previous_rules)

    if resolved:
        _styled_run(
            document.add_paragraph(), f"已解決（{len(resolved)} 項）：", bold=True,
        )
        for rule in resolved:
            _styled_run(
                document.add_paragraph(style="List Bullet"),
                previous_titles.get(rule, rule),
            )
    if appeared:
        _styled_run(
            document.add_paragraph(), f"新出現（{len(appeared)} 項）：", bold=True,
        )
        for rule in appeared:
            _styled_run(document.add_paragraph(style="List Bullet"), current_map[rule])
    if not resolved and not appeared:
        _styled_run(document.add_paragraph(), "發現項目與前次相同，沒有新增或解決。")


def _add_page_list(document, scan_job: ScanJob, counter) -> None:
    """掃描頁面清單。讓收件者知道報告到底看過哪些頁面。"""
    pages = list(
        scan_job.pages.all()
        .only("url", "final_url", "title", "status_code", "depth", "blocked_reason")
        .order_by("depth", "url")
    )
    if not pages:
        return
    _heading(document, f"{counter.next()}　掃描頁面清單", level=2)
    _styled_run(
        document.add_paragraph(),
        f"本次共擷取 {len(pages)} 個頁面。標示「已略過」的頁面未納入分析。",
    )
    table = document.add_table(rows=len(pages) + 1, cols=4)
    table.style = "Table Grid"
    _header_row(table, ["網址", "頁面標題", "狀態", "深度"])
    for index, page in enumerate(pages, start=1):
        cells = table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], page.final_url or page.url, size=8)
        _styled_run(cells[1].paragraphs[0], page.title or "（無標題）", size=8)
        status = str(page.status_code or "—")
        if page.blocked_reason:
            status = f"{status}　已略過：{page.blocked_reason}"
        _styled_run(cells[2].paragraphs[0], status, size=8)
        _styled_run(cells[3].paragraphs[0], str(page.depth), size=8)


def _add_scan_scope(document, scan_job: ScanJob) -> None:
    """掃描範圍。收件者要能判斷這份報告涵蓋了什麼，「沒發現問題」才有意義。

    scope 一律取自 scan_plan.build_scan_execution_plan()，不要在這裡重複
    「max_pages==1 代表單頁」的慣例，避免兩邊漂移。
    """
    _heading(document, "2　掃描範圍", level=1)
    document.add_paragraph()
    scope_label = "單頁" if build_scan_execution_plan(scan_job).scope == "single" else "全網站"
    _kv_table(document, [
        ("掃描範圍", scope_label),
        ("探測模式", scan_job.get_scan_mode_display()),
        ("頁數上限", str(scan_job.max_pages)),
        ("連結深度上限", str(scan_job.max_depth)),
        ("實際掃描頁數", str(scan_job.pages.count())),
        ("遵守 robots.txt", "是" if scan_job.respect_robots else "否"),
    ])
    _add_entry_screenshot(document, scan_job)
    _add_scan_warnings(document, scan_job)
    document.add_page_break()


def _add_entry_screenshot(document, scan_job: ScanJob) -> None:
    """入口頁截圖。

    刻意只放這一張：全頁截圖體積大，把 50 頁的掃描全塞進去會讓 .docx 失控，
    而 header / DNS / meta 這類發現本來就沒有視覺佐證價值。放一張的用途是讓
    收件者一眼確認「這份報告講的確實是我的網站」。

    路徑慣例與 views.py 的 screenshot action 一致：screenshot_path 相對於
    BASE_DIR。檔案不在（保留期限、換機器、S3 模式）就整段略過，不影響報告。
    """
    entry_page = (
        scan_job.pages.exclude(screenshot_path="")
        .order_by("depth", "id")
        .first()
    )
    if entry_page is None:
        return
    screenshot = Path(settings.BASE_DIR) / entry_page.screenshot_path
    if not screenshot.exists():
        return
    _heading(document, "掃描當下的網站畫面", level=2)
    try:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(screenshot), width=Inches(5.2))
    except Exception:  # noqa: BLE001 — 截圖損毀不該讓整份報告產不出來
        return
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _styled_run(
        caption,
        f"{entry_page.final_url or entry_page.url}（掃描當下擷取）",
        size=9, color=ARGUS_MUTED,
    )


def _add_scan_warnings(document, scan_job: ScanJob) -> None:
    """掃描過程的警示。

    只挑對收件者有意義的項目。內部運維資訊（settlement_error 的計費結算、
    agent 的 token 用量）刻意不寫進對外報告。
    """
    warnings = scan_job.warning_summary or {}
    lines: list[str] = []

    if warnings.get("scan_effectiveness") == "no_pages_crawled":
        lines.append(
            "掃描有效性警示：本次未抓到任何頁面（目標可能不可達或全部逾時）。"
            "SEO 與 AEO 未評估，分數僅反映站台層級檢查，不應解讀為「網站沒有問題」。"
        )

    blocked = warnings.get("blocked_urls") or []
    if isinstance(blocked, list) and blocked:
        lines.append(f"依 robots.txt 或掃描範圍限制，略過 {len(blocked)} 個頁面未檢查。")

    failed = warnings.get("failed_urls") or []
    if isinstance(failed, list) and failed:
        lines.append(f"有 {len(failed)} 個頁面擷取失敗（逾時或回應異常），未納入本次分析。")

    tech_stack = warnings.get("tech_stack") or []
    if isinstance(tech_stack, list) and tech_stack:
        lines.append(f"偵測到的技術棧：{'、'.join(str(item) for item in tech_stack)}")

    if not lines:
        return
    _heading(document, "掃描警示", level=2)
    for line in lines:
        _styled_run(document.add_paragraph(style="List Bullet"), line)


def _add_top_actions(document, scan_job: ScanJob) -> None:
    _heading(document, "3　優先改善建議", level=1)
    document.add_paragraph()
    actions = scan_job.top_actions or []
    if not actions:
        _styled_run(
            document.add_paragraph(),
            "本次掃描沒有需要優先處理的項目。",
        )
        document.add_page_break()
        return
    _styled_run(
        document.add_paragraph(),
        "以下是本次掃描中最值得先處理的項目，依影響程度排序。"
        "詳細說明請見第 4 章對應的發現項目。",
    )
    table = document.add_table(rows=len(actions) + 1, cols=3)
    table.style = "Table Grid"
    _header_row(table, ["順序", "嚴重度", "問題"])
    for index, action in enumerate(actions, start=1):
        cells = table.rows[index].cells
        severity = action.get("severity", "")
        _styled_run(cells[0].paragraphs[0], str(index), bold=True)
        _styled_run(
            cells[1].paragraphs[0], get_severity_display(severity),
            bold=True, color=SEVERITY_COLOR.get(severity, "000000"),
        )
        _styled_run(
            cells[2].paragraphs[0],
            f"{action.get('title', '')}"
            f"（{CATEGORY_DISPLAY.get(action.get('category', ''), '')}）",
        )
    document.add_page_break()


def _description_for_report(finding) -> str:
    """去掉 description 開頭與報告行為矛盾的警語。

    scanners.py 的 PII finding 在 description 開頭寫「⚠️ 此項目顯示原始個資，
    請依個資法妥善處理本報告。」——那句對前端成立（依使用者要求，API/畫面顯示
    未遮罩的 evidence），但報告會遮罩（09******90），照搬進來就是假話。報告本來
    就會在「檢測依據」下輸出自己那句正確的遮罩提示，不需要這一句。

    規則故意寫得寬鬆（剝掉開頭所有 ⚠️ 起始行）而不是比對特定字串：警語措辭改了
    也不會漏掉，而 description 的實質內容不會以 ⚠️ 開頭。
    """
    lines = (finding.description or "").splitlines()
    while lines and lines[0].lstrip().startswith("⚠️"):
        lines.pop(0)
    return "\n".join(lines).strip() or "（無）"


def _add_findings(document, grouped_findings) -> None:
    _heading(document, "4　發現項目", level=1)
    if not grouped_findings:
        _styled_run(document.add_paragraph(), "本次掃描沒有發現需要記錄的項目。")
        document.add_page_break()
        return

    for index, item in enumerate(grouped_findings, start=1):
        finding = item["finding"]
        pages = item["pages"]
        severity_color = SEVERITY_COLOR.get(finding.severity, "000000")

        heading = document.add_heading(level=2)
        _styled_run(heading, f"4.{index}　{finding.title}", bold=True, color=severity_color)

        _kv_table(document, [
            ("嚴重度", get_severity_display(finding.severity)),
            ("分類", CATEGORY_DISPLAY.get(finding.category, finding.category.upper())),
            (
                "受影響頁面",
                f"共 {len(pages)} 處" if len(pages) > 1 else pages[0],
            ),
        ])
        if len(pages) > 1:
            for page_url in pages:
                _styled_run(document.add_paragraph(style="List Bullet"), page_url, size=9)

        _heading(document, "問題是什麼", level=3)
        _styled_run(document.add_paragraph(), _description_for_report(finding))

        is_info = finding.severity == Finding.Severity.INFO
        if is_info:
            _heading(document, "這代表什麼", level=3)
            _styled_run(document.add_paragraph(), INFO_NOTE)
        else:
            _heading(document, "為什麼要在意", level=3)
            impact = CATEGORY_IMPACT.get(finding.category, "請依你的業務情境評估影響。")
            urgency = SEVERITY_URGENCY.get(finding.severity, "")
            _styled_run(document.add_paragraph(), f"{impact}{urgency}")

        _heading(document, "怎麼修", level=3)
        _styled_run(document.add_paragraph(), finding.remediation or "（無）")

        if not is_info:
            # info 項目不會在下次掃描「消失」（正向指標本來就該留著，小問題也未必
            # 值得專程複驗），叫讀者去確認它不見了只會造成困惑。
            _heading(document, "修好了怎麼確認", level=3)
            _styled_run(
                document.add_paragraph(),
                CATEGORY_VERIFY.get(finding.category, "修補後重新執行一次 Argus 掃描確認。"),
            )

        if finding.evidence:
            _heading(document, "檢測依據", level=3)
            # 先在完整字串上遮罩、再截斷顯示長度：反過來做的話，PII 數值可能剛好被
            # 截斷點切一半，殘缺數字長度不足以命中 regex，反而以明文殘留。
            masked_evidence = mask_pii_evidence(finding.evidence)
            if masked_evidence != finding.evidence:
                _styled_run(
                    document.add_paragraph(),
                    "⚠️ 以下內容為偵測到之敏感樣本部分遮罩後的結果，"
                    "請依個資法妥善保管本報告。",
                    bold=True, color=SEVERITY_COLOR["high"],
                )
            evidence_text = masked_evidence[:1000]
            truncated = "…（證據過長已截斷）" if len(masked_evidence) > 1000 else ""
            _styled_run(document.add_paragraph(), f"{evidence_text}{truncated}", size=9)

        if finding.ai_handoff_prompt:
            # 報告不做 AI 解釋（ai_explanation 從未被實作），但每筆 finding 都已經有
            # 一段組好的提示詞，使用者可以直接貼進 ChatGPT / Claude 取得深入說明。
            # 必須跟 evidence 一樣遮罩：build_ai_handoff_prompt() 內嵌了原始 evidence，
            # 不遮罩等於從後門把個資漏回這份會被轉寄的報告。
            _heading(document, "想更深入了解？", level=3)
            _styled_run(
                document.add_paragraph(),
                "把下面這段文字複製貼給 ChatGPT、Claude 等 AI 助手，可取得更詳細的說明：",
            )
            masked_prompt = mask_pii_evidence(finding.ai_handoff_prompt)
            prompt_text = masked_prompt[:1500]
            prompt_note = "…（提示詞過長已截斷）" if len(masked_prompt) > 1500 else ""
            _styled_run(
                document.add_paragraph(), f"{prompt_text}{prompt_note}",
                size=9, color=ARGUS_MUTED,
            )

        document.add_paragraph()
    document.add_page_break()


def _add_authorization(document, scan_job: ScanJob, section_number: str) -> None:
    """掃描授權聲明。

    Argus 是授權式掃描平台，一份對外的資安報告必須記載這次掃描的授權依據，
    否則等於放棄本專案最核心的合規主張。

    刻意不寫入 AuthorizationConsent 的 ip_address 與 user_agent，也不寫授權
    帳號：這份 .docx 會被下載、轉寄、存檔給第三方，授權人的 IP 與瀏覽器指紋
    是個資，寫進去沒有任何對收件者的價值，只增加外洩面。需要稽核時查 DB 與
    AdminAuditLog。
    """
    _heading(document, f"{section_number}　掃描授權聲明", level=2)
    consent = getattr(scan_job, "authorization_consent", None)
    if consent is None:
        _styled_run(
            document.add_paragraph(),
            "查無授權紀錄：本次掃描在資料庫中沒有對應的授權同意紀錄。"
            "若這份報告要作為稽核依據，請先確認授權來源。",
        )
        return
    _kv_table(document, [
        ("授權網域", consent.authorized_domain),
        ("授權時間", consent.created_at.strftime("%Y-%m-%d %H:%M:%S")),
        (
            "主動測試授權",
            "是" if consent.active_testing_authorized else "否（僅被動偵測）",
        ),
    ])
    _styled_run(document.add_paragraph(), "授權聲明內容：", bold=True)
    _styled_run(document.add_paragraph(), consent.statement)


def _collect_glossary_terms(grouped_findings) -> list[tuple[str, str]]:
    """只挑這份報告裡真的出現過的術語，不是貼一份固定清單。"""
    corpus_parts: list[str] = []
    for item in grouped_findings:
        finding = item["finding"]
        corpus_parts += [
            finding.title or "", finding.description or "",
            finding.remediation or "", finding.evidence or "",
            finding.owasp_category or "", finding.cwe_id or "",
        ]
    corpus = "\n".join(corpus_parts).lower()
    return [
        (term, explanation)
        for term, explanation in GLOSSARY.items()
        if term.lower() in corpus
    ]


def _add_glossary(document, grouped_findings, counter) -> None:
    terms = _collect_glossary_terms(grouped_findings)
    if not terms:
        return
    _heading(document, f"{counter.next()}　名詞解釋", level=2)
    _styled_run(
        document.add_paragraph(),
        "以下是本報告中出現的專有名詞。只列出這次真的用到的，方便你對照閱讀。",
    )
    table = document.add_table(rows=len(terms) + 1, cols=2)
    table.style = "Table Grid"
    _header_row(table, ["名詞", "白話說明"])
    for index, (term, explanation) in enumerate(terms, start=1):
        cells = table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], term, bold=True)
        _styled_run(cells[1].paragraphs[0], explanation)


def _add_technical_index(document, grouped_findings, counter) -> None:
    """技術索引：rule_id / OWASP / CWE 收在這裡，不混進正文。

    這些識別碼對網站主沒有意義，但工程師與稽核需要，所以保留但下放到附錄。
    """
    if not grouped_findings:
        return
    _heading(document, f"{counter.next()}　技術索引", level=2)
    _styled_run(
        document.add_paragraph(),
        "供工程師與稽核人員對照使用。OWASP 是國際公認的網站安全風險分類，"
        "CWE 是通用的軟體弱點編號；一般讀者可略過本節。",
    )
    table = document.add_table(rows=len(grouped_findings) + 1, cols=4)
    table.style = "Table Grid"
    _header_row(table, ["項次", "問題", "規則 ID", "OWASP / CWE"])
    for index, item in enumerate(grouped_findings, start=1):
        finding = item["finding"]
        cells = table.rows[index].cells
        _styled_run(cells[0].paragraphs[0], f"4.{index}", size=9)
        _styled_run(cells[1].paragraphs[0], finding.title, size=9)
        _styled_run(cells[2].paragraphs[0], finding.rule_id or "—", size=9)
        owasp = finding.owasp_category or "—"
        cwe = finding.cwe_id or "—"
        _styled_run(cells[3].paragraphs[0], f"{owasp} / {cwe}", size=9)


def _add_appendix(document, scan_job: ScanJob, grouped_findings) -> None:
    _heading(document, "5　附錄", level=1)
    # 小節編號動態產生：名詞解釋與技術索引在沒有 finding 時會整段消失，頁面清單
    # 在爬 0 頁時也不出現，寫死號碼會出現「5.2 之後接 5.4」這種跳號。
    counter = _SectionNumber("5")
    _add_authorization(document, scan_job, counter.next())
    _add_glossary(document, grouped_findings, counter)
    _add_technical_index(document, grouped_findings, counter)
    _add_page_list(document, scan_job, counter)

    _heading(document, f"{counter.next()}　免責聲明", level=2)
    _styled_run(document.add_paragraph(), DISCLAIMER)

    _heading(document, f"{counter.next()}　本報告如何產生", level=2)
    # 舊版寫「再交由 AI 進行自然語言解釋與改善建議撰寫」，但 Finding.ai_explanation
    # 與 ai_remediation 在整個 backend 只被寫入空字串，該功能從未實作——這是一份
    # 對外交付文件裡的不實陳述。改為只描述實際做到的事。
    _styled_run(
        document.add_paragraph(),
        "本報告採 Evidence-first 原則：SEO、AEO、GEO 與資安判斷全部來自爬蟲與規則引擎"
        "產生的可驗證證據，每一項發現都附上當下實際觀測到的內容。"
        "報告產生過程不使用 AI 改寫或推論結論，因此不會出現無法追溯到掃描證據的說法。"
        "每一項發現另附一段可直接貼給 AI 助手的提示詞，方便你取得更深入的說明。",
    )


def report_output_path(scan_job: ScanJob) -> Path:
    """報告檔案位置。views.py 判斷快取時也用這支，避免兩邊各寫一次檔名慣例。"""
    return Path(settings.MEDIA_ROOT) / "reports" / f"scan-{scan_job.id}-report.docx"


def build_scan_report(scan_job: ScanJob) -> str:
    output_path = report_output_path(scan_job)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    grouped_findings = _group_findings_for_report(
        scan_job.findings.select_related("page").all()
    )

    report_number = build_report_number(scan_job)
    generated_at = timezone.now()

    document = Document()
    _add_header_footer(document, scan_job, report_number)
    _add_cover(document, scan_job, report_number)
    _add_table_of_contents(document, grouped_findings)
    _add_summary(document, scan_job, grouped_findings)
    _add_scan_scope(document, scan_job)
    _add_top_actions(document, scan_job)
    _add_findings(document, grouped_findings)
    _add_appendix(document, scan_job, grouped_findings)
    document.save(output_path)

    # 雜湊算完才寫防偽紀錄：報告本身不印雜湊（印了就循環相依——雜湊要涵蓋整份
    # 檔案，而檔案裡又要有雜湊），收件者拿編號到查驗頁取得雜湊自行比對。
    ReportVerification.objects.update_or_create(
        scan_job=scan_job,
        defaults={
            "report_number": report_number,
            "content_sha256": sha256(output_path.read_bytes()).hexdigest(),
            "generated_at": generated_at,
        },
    )
    return str(output_path)
