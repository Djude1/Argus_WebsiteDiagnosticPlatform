"""報告防偽與快取測試（稽核第四階段：E2、E3、E4、G1）。

背景（docs/scan-report-quality-audit-2026-08-30.md）：
scan 25 的報告全文只出現 1 次「Argus」，沒有報告編號、沒有產生時間、沒有任何
驗證機制——任何人都能用 Word 改掉分數再轉寄，宣稱是 Argus 出具的報告。
且 views.py 每次下載都重新產生一份，同一個掃描會產出內容不同的多份檔案，
連「哪一份才是正本」都無從分辨。
"""

from __future__ import annotations

from pathlib import Path
from unittest import mock

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.urls import reverse
from docx import Document
from rest_framework.test import APIClient

from apps.scans.models import Finding, ReportVerification, ScanJob
from apps.scans.report_render import RENDERER_VERSION
from apps.scans.reports import build_scan_report, report_output_path

User = get_user_model()


class ReportVerificationTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="verify", password="safe-test-password")
        self.scan_job = ScanJob.objects.create(
            user=self.user,
            original_url="https://example.com/",
            normalized_url="https://example.com/",
            origin="example.com",
            status=ScanJob.Status.COMPLETED,
            overall_score=65,
            category_scores={"security": 15},
        )
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="high",
            category=Finding.Category.SECURITY, title="缺少 HSTS",
            description="d", remediation="r", rule_id="hsts",
            ai_handoff_prompt="p", priority_score=75.0,
        )

    def _text(self) -> str:
        document = Document(build_scan_report(self.scan_job))
        parts = [p.text for p in document.paragraphs]
        parts += [p.text for p in document.sections[0].footer.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    # --- E2 報告編號與產生時間 ---------------------------------------
    def test_building_a_report_records_a_verification_row(self):
        build_scan_report(self.scan_job)

        record = ReportVerification.objects.get(scan_job=self.scan_job)
        self.assertTrue(record.report_number.startswith("ARGUS-"))
        self.assertEqual(len(record.content_sha256), 64)

    def test_report_number_is_stable_across_regeneration(self):
        """既有報告已經流出去了，重新產生不能讓舊編號失效。"""
        build_scan_report(self.scan_job)
        first = ReportVerification.objects.get(scan_job=self.scan_job).report_number

        build_scan_report(self.scan_job)
        second = ReportVerification.objects.get(scan_job=self.scan_job).report_number

        self.assertEqual(first, second)
        self.assertEqual(ReportVerification.objects.filter(scan_job=self.scan_job).count(), 1)

    def test_report_number_appears_on_cover_and_footer(self):
        build_scan_report(self.scan_job)
        number = ReportVerification.objects.get(scan_job=self.scan_job).report_number

        self.assertGreaterEqual(self._text().count(number), 2)

    def test_content_hash_matches_the_saved_file(self):
        from hashlib import sha256

        path = Path(build_scan_report(self.scan_job))
        record = ReportVerification.objects.get(scan_job=self.scan_job)

        self.assertEqual(sha256(path.read_bytes()).hexdigest(), record.content_sha256)

    # --- E4 免責聲明 -------------------------------------------------
    def test_report_carries_a_disclaimer(self):
        text = self._text()
        self.assertIn("免責", text)
        self.assertIn("滲透測試", text)

    # --- E3 公開查驗端點 ---------------------------------------------
    def test_verify_endpoint_is_public_and_returns_scan_facts(self):
        build_scan_report(self.scan_job)
        number = ReportVerification.objects.get(scan_job=self.scan_job).report_number

        response = APIClient().get(reverse("report-verify", args=[number]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.data["report_number"], number)
        self.assertEqual(response.data["scan_target"], "https://example.com/")
        self.assertEqual(response.data["overall_score"], 65)

    def test_verify_endpoint_never_exposes_who_ran_the_scan(self):
        """查驗頁是公開的，不能靠報告編號反查出使用者身分。"""
        build_scan_report(self.scan_job)
        number = ReportVerification.objects.get(scan_job=self.scan_job).report_number

        payload = str(APIClient().get(reverse("report-verify", args=[number])).data)

        self.assertNotIn(self.user.username, payload)
        self.assertNotIn("user", payload.lower().replace("username", ""))

    def test_unknown_report_number_returns_404(self):
        response = APIClient().get(reverse("report-verify", args=["ARGUS-9-20260101-DEAD"]))

        self.assertEqual(response.status_code, 404)


class ReportCacheTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="cache", password="safe-test-password")
        self.scan_job = ScanJob.objects.create(
            user=self.user, original_url="https://example.com/",
            normalized_url="https://example.com/", origin="example.com",
            status=ScanJob.Status.COMPLETED, overall_score=65,
        )
        self.client = APIClient()
        self.client.force_authenticate(user=self.user)

    def test_second_download_does_not_regenerate_the_report(self):
        """重複產生會改變內容指紋，讓已交付出去的副本在查驗頁對不上。

        斷言用「產生器被呼叫幾次」而不是比對指紋：兩次下載若落在同一秒，
        報告裡的產生時間字串相同，指紋可能碰巧一致而讓測試僥倖通過。
        """
        first = self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        self.assertEqual(first.status_code, 200)

        with mock.patch("apps.scans.views.build_scan_report") as build:
            second = self.client.get(f"/api/scans/{self.scan_job.id}/report/")

        self.assertEqual(second.status_code, 200)
        build.assert_not_called()

    def test_missing_file_is_rebuilt_even_if_a_verification_row_exists(self):
        """磁碟上的報告被清掉（保留期限、換機器）時仍要能重新產生。"""
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        path = report_output_path(self.scan_job)
        path.unlink()

        response = self.client.get(f"/api/scans/{self.scan_job.id}/report/")

        self.assertEqual(response.status_code, 200)
        self.assertTrue(path.exists())

    def test_stale_renderer_version_forces_a_rebuild(self):
        """排版升級後，既有掃描下載到的必須是新版面。

        沒有這道判斷時，掃描一旦產過報告就永遠鎖在舊排版——使用者實際踩過：
        修好圖表後重新下載舊掃描的報告，拿到的是沒有圖表的快取檔。
        """
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        ReportVerification.objects.filter(scan_job=self.scan_job).update(
            renderer_version=0
        )

        with mock.patch(
            "apps.scans.views.build_scan_report",
            wraps=build_scan_report,
        ) as build:
            response = self.client.get(f"/api/scans/{self.scan_job.id}/report/")

        self.assertEqual(response.status_code, 200)
        build.assert_called_once()
        record = ReportVerification.objects.get(scan_job=self.scan_job)
        self.assertEqual(record.renderer_version, RENDERER_VERSION)

    def test_rebuild_keeps_the_previous_hash_verifiable(self):
        """重產換掉指紋，但先前交付出去的副本不能因此被判成偽造。"""
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        delivered = ReportVerification.objects.get(scan_job=self.scan_job).content_sha256

        # 換掉檔案內容，確保重產一定產生不同指紋（同一秒內重產可能位元組相同）
        report_output_path(self.scan_job).write_bytes(b"tampered")
        ReportVerification.objects.filter(scan_job=self.scan_job).update(
            renderer_version=0, content_sha256=delivered
        )
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")

        record = ReportVerification.objects.get(scan_job=self.scan_job)
        self.assertNotEqual(record.content_sha256, delivered)
        self.assertIn(delivered, record.previous_sha256)

    def test_verify_endpoint_accepts_a_superseded_hash(self):
        """收件者拿舊版報告來查驗，仍要回答「是真的」。"""
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        record = ReportVerification.objects.get(scan_job=self.scan_job)
        old_hash = "a" * 64
        record.previous_sha256 = [old_hash]
        record.save(update_fields=["previous_sha256"])

        url = reverse("report-verify", args=[record.report_number])
        response = APIClient().get(url, {"content_sha256": old_hash})

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.data["matches"])
        self.assertFalse(response.data["is_latest_version"])

    def test_verify_endpoint_rejects_an_unknown_hash(self):
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        record = ReportVerification.objects.get(scan_job=self.scan_job)

        url = reverse("report-verify", args=[record.report_number])
        response = APIClient().get(url, {"content_sha256": "b" * 64})

        self.assertEqual(response.status_code, 200)
        self.assertFalse(response.data["matches"])

    def test_verify_endpoint_never_lists_historical_hashes(self):
        """揭露一整串指紋對驗證沒有幫助，只會讓人以為要逐一比對。"""
        self.client.get(f"/api/scans/{self.scan_job.id}/report/")
        record = ReportVerification.objects.get(scan_job=self.scan_job)
        record.previous_sha256 = ["c" * 64]
        record.save(update_fields=["previous_sha256"])

        url = reverse("report-verify", args=[record.report_number])
        response = APIClient().get(url)

        self.assertNotIn("previous_sha256", response.data)
        self.assertNotIn("c" * 64, str(response.data))
