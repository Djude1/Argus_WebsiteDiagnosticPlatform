"""Argus 網站健檢報告 generator — data (JSON) in, styled .docx out.

Public API:
    generate_report(data: dict, out_path: str, workdir: str = None) -> str

The design (layout, palette, cards, watermark) is fixed here; the Agent supplies
only the data. See schema.json for the expected input shape.
"""
import os
import tempfile

from docx import Document
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import theme as T
from . import charts as C
from . import _xml as X

EMU_PER_IN = 914400


# ---------- run / paragraph helpers ----------
def _rgb(hexstr):
    return RGBColor.from_string(hexstr)


def _set_font(run, mono=False):
    name = T.FONT_MONO if mono else T.FONT_CJK
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(a), name)


def add_run(paragraph, text, size=10, color=T.SLATE, bold=False, mono=False, chip=None):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = _rgb(color)
    _set_font(run, mono=mono)
    if chip:
        # shaded run (severity chip)
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), chip)
        rPr.append(shd)
    return run


def add_para(doc_or_cell, runs=None, align=None, before=0, after=6, line=1.15,
             keep_next=False):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    if align:
        p.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if keep_next:
        pf.keep_with_next = True
    if runs:
        for r in runs:
            add_run(p, **r)
    return p


def chip_run(paragraph, severity):
    meta = T.SEVERITY[severity]
    add_run(paragraph, f" {severity} ", size=9, color=meta["text"], bold=True,
            chip=meta["fill"])


# ---------- headings ----------
def h1(doc, num, title, page_break=True):
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(0 if page_break else 18)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, f"{num}　", size=15, color=T.NAVY, bold=True)
    add_run(p, title, size=15, color=T.NAVY, bold=True)
    X.set_para_borders(p, {"bottom": (18, T.NAVY, "6")})
    return p


def h2(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    add_run(p, title, size=12, color=T.NAVY, bold=True)
    X.set_para_borders(p, {"left": (24, T.NAVY, "8")})
    X.set_para_indent(p, left=170)
    return p


# ---------- tables ----------
def _dxa(cols):  # column widths list
    return cols


def data_table(doc, headers, rows, widths):
    """rows: list of lists; each cell is str OR list-of-run-dicts OR ('chip',sev)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    # header
    hdr = table.rows[0]
    X.set_row_cant_split(hdr)
    hdr._tr.get_or_add_trPr().append(_tbl_header_mark())
    for i, htext in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Emu(int(widths[i] * 635))
        _fill_cell(cell, [[{"text": htext, "size": 9.5, "color": T.NAVY, "bold": True}]],
                   fill=T.BGBLUE)
    # body
    for ri, row in enumerate(rows):
        tr = table.add_row()
        X.set_row_cant_split(tr)
        fill = T.BG if ri % 2 == 1 else None
        for i, cell_data in enumerate(row):
            cell = tr.cells[i]
            cell.width = Emu(int(widths[i] * 635))
            _fill_cell(cell, _normalize_cell(cell_data), fill=fill)
    _apply_table_borders(table)
    return table


def _normalize_cell(cell_data):
    """Return list-of-paragraphs, each a list-of-run-dicts."""
    if isinstance(cell_data, tuple) and cell_data[0] == "chip":
        return [("chip", cell_data[1])]
    if isinstance(cell_data, list):
        return [cell_data]  # single paragraph of runs
    return [[{"text": str(cell_data), "size": 9.5, "color": T.SLATE}]]


def _fill_cell(cell, paragraphs, fill=None):
    # clear default empty paragraph
    cell.text = ""
    first = True
    for para in paragraphs:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        if isinstance(para, tuple) and para[0] == "chip":
            chip_run(p, para[1])
        else:
            for r in para:
                add_run(p, **r)
    if fill:
        X.set_cell_shading(cell, fill)
    X.set_cell_margins(cell)
    X.set_cell_vertical_center(cell)
    X.set_cell_borders(cell, sz=2, color=T.LINE)


def _tbl_header_mark():
    e = OxmlElement("w:tblHeader")
    return e


def _apply_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, color in (("top", T.BORD), ("left", T.BORD), ("bottom", T.BORD),
                        ("right", T.BORD), ("insideH", T.LINE), ("insideV", T.LINE)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "2")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


# ---------- images ----------
def add_image(doc_or_cell, path, width_in):
    p = doc_or_cell.add_paragraph() if hasattr(doc_or_cell, "add_paragraph") else doc_or_cell
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Emu(int(width_in * EMU_PER_IN)))
    return p


# ---------- finding card ----------
def finding_card(doc, f):
    # header band
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    add_run(p, f'{f["id"]}　', size=11, color=T.NAVY, bold=True)
    add_run(p, f["title"], size=11, color=T.NAVY, bold=True)
    X.set_para_shading(p, T.BGBLUE)
    X.set_para_borders(p, {"top": (4, T.NAVY, "0"), "left": (24, T.NAVY, "8"),
                           "bottom": (4, T.NAVY, "0"), "right": (4, T.NAVY, "0")})
    X.set_para_indent(p, left=100, right=80)
    # meta line
    mp = doc.add_paragraph()
    mp.paragraph_format.space_before = Pt(3)
    mp.paragraph_format.space_after = Pt(3)
    mp.paragraph_format.keep_with_next = True
    chip_run(mp, f["severity"])
    add_run(mp, f'　·　{f["category"]}　·　{f["scope"]}', size=9, color=T.GREY)
    # affected urls
    if f.get("urls"):
        up = doc.add_paragraph()
        up.paragraph_format.space_after = Pt(4)
        add_run(up, f["urls"], size=8, color=T.LIGHTGREY)
    # problem
    _card_label(doc, "▍問題是什麼", T.SLATE)
    add_para(doc, [{"text": f["problem"], "size": 10, "color": T.SLATE}],
             after=4, line=1.4)
    # fix (action zone)
    _card_label(doc, "▍怎麼修", T.NAVY)
    fp = add_para(doc, [{"text": f["fix"], "size": 10, "color": T.SLATE}],
                  after=4, line=1.4)
    X.set_para_shading(fp, T.FIXBLUE)
    X.set_para_indent(fp, left=130, right=130)
    X.set_para_borders(fp, {"left": (18, "0369A1", "4")})
    # evidence
    _card_label(doc, "▍檢測依據", T.GREY)
    ep = add_para(doc, [{"text": f["evidence"], "size": 9, "color": T.SLATE, "mono": True}],
                  before=2, after=3, line=1.35)
    X.set_para_shading(ep, T.BG)
    X.set_para_indent(ep, left=130, right=130)
    X.set_para_borders(ep, {"left": (18, T.LIGHTGREY, "4")})


def _card_label(doc, text, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    add_run(p, text, size=9, color=color, bold=True)


def severity_group_header(doc, severity, count):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    chip_run(p, severity)
    add_run(p, f"  {severity}項目（{count} 項）", size=11, color=T.NAVY, bold=True)
    X.set_para_borders(p, {"left": (24, T.SEVERITY[severity]["fill"], "8")})
    X.set_para_indent(p, left=170)


# ---------- section/header/footer ----------
def _setup_section(doc):
    sec = doc.sections[0]
    sec.page_width = Emu(12240 * 635)
    sec.page_height = Emu(15840 * 635)
    sec.top_margin = Emu(1440 * 635)
    sec.bottom_margin = Emu(1440 * 635)
    sec.left_margin = Emu(1700 * 635)
    sec.right_margin = Emu(1700 * 635)
    sec.header_distance = Emu(720 * 635)
    sec.footer_distance = Emu(600 * 635)
    return sec


def _build_header(sec, site_url):
    header = sec.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"ARGUS 網站健檢報告　|　{site_url}", size=9, color="5B6B7C")
    X.add_watermark_to_header(header)


def _build_footer(sec, report_id):
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    X.set_para_borders(p, {"top": (4, T.LINE, "6")})
    add_run(p, "Argus 授權式網站健檢　|　第 ", size=8, color=T.GREY)
    X.add_page_number_field(p)
    add_run(p, f" 頁　|　{report_id}", size=8, color=T.GREY)


# ---------- main ----------
def generate_report(data, out_path, workdir=None):
    workdir = workdir or tempfile.mkdtemp(prefix="argus_")
    chart_paths = C.build_all(data, workdir)

    doc = Document()
    # default font
    style = doc.styles["Normal"]
    style.font.name = T.FONT_CJK
    style.font.size = Pt(10)
    style.font.color.rgb = _rgb(T.SLATE)
    _set_font_style_eastasia(style)

    sec = _setup_section(doc)
    s = data["summary"]
    meta = data["meta"]
    _build_header(sec, meta["site_url"])
    _build_footer(sec, meta["report_id"])

    _cover(doc, data, chart_paths)
    _summary(doc, data, chart_paths)
    _priorities(doc, data)
    _why_matters(doc, data)
    _findings(doc, data)
    _scan_info(doc, data, chart_paths)
    _appendix(doc, data)

    X.fix_zoom_setting(doc)
    doc.save(out_path)
    return out_path


def _set_font_style_eastasia(style):
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), T.FONT_CJK)


# ---------- sections ----------
def _cover(doc, data, ch):
    s, meta = data["summary"], data["meta"]
    add_para(doc, before=70, after=0)
    add_para(doc, [{"text": "ARGUS", "size": 38, "color": T.NAVY, "bold": True}],
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, [{"text": "網站健檢報告", "size": 18, "color": T.SLATE}],
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_image(doc, ch["score"], 2.0)
    band = T.score_band_label(s["overall_score"])
    add_para(doc, [
        {"text": f'整體分數 {s["overall_score"]} / 100　', "size": 13,
         "color": T.score_band_color(s["overall_score"]), "bold": True},
        {"text": f"（{band}）", "size": 12, "color": T.GREY}],
        align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3)
    add_para(doc, [{"text": meta["site_url"], "size": 11, "color": T.NAVY, "bold": True}],
             align=WD_ALIGN_PARAGRAPH.CENTER, after=15)
    add_para(doc, [{"text": f'掃描完成：{s["scan_date"]}　|　報告產生：{meta.get("generated_at", s["scan_date"])}',
                    "size": 9, "color": T.GREY}],
             align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    add_para(doc, [{"text": "報告編號　", "size": 10, "color": T.SLATE},
                   {"text": meta["report_id"], "size": 10, "color": T.NAVY, "bold": True}],
             align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=1)
    add_para(doc, [{"text": "可至 Argus 網站「報告查驗」頁（/verify）輸入編號，核對本報告真偽與內容指紋。",
                    "size": 8, "color": T.LIGHTGREY}],
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def _summary(doc, data, ch):
    s = data["summary"]
    h1(doc, "1", "一頁摘要")
    if s.get("headline"):
        add_para(doc, [{"text": s["headline"], "size": 10, "color": T.SLATE}],
                 after=8, line=1.44)
    h2(doc, "各分類分數")
    add_image(doc, ch["categories"], 5.5)
    add_para(doc, [{"text": "虛線為 60 / 80 分門檻。分數為各「已評估」分類的平均；標示「未評估」者不納入計算。",
                    "size": 8, "color": T.LIGHTGREY}], after=8)
    h2(doc, "發現項目分佈")
    add_image(doc, ch["severity"], 5.5)
    counts = {}
    for f in data["findings"]:
        counts[f["severity"]] = counts.get(f["severity"], 0) + 1
    runs = []
    for sev in T.SEVERITY_ORDER:
        if counts.get(sev):
            runs.append(("chip", sev))
            runs.append({"text": f"  {counts[sev]} 項　", "size": 9, "color": T.GREY})
    # chips need special handling
    p = add_para(doc, after=8)
    for r in runs:
        if isinstance(r, tuple):
            chip_run(p, r[1])
        else:
            add_run(p, **r)
    # trend
    if s.get("previous") and "trend" in ch:
        h2(doc, "與前次掃描比較")
        _trend_block(doc, s, ch)
    # scoring legend (fixed reference)
    h2(doc, "分數怎麼看")
    data_table(doc, ["分數", "評級", "建議"], [
        ["80–100", "良好", "持續維持即可，建議定期複檢。"],
        ["60–79", "需改善", "有幾項體質問題值得排入維護排程。"],
        ["40–59", "建議儘快處理", "累積的問題已可能影響流量或安全，建議近期處理。"],
        ["0–39", "需優先處理", "存在較高風險的項目，建議優先安排修補。"],
    ], [1600, 2100, 4400])


def _trend_block(doc, s, ch):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left, right = table.rows[0].cells
    left.width = Emu(int(3400 * 635))
    right.width = Emu(int(4700 * 635))
    for c in (left, right):
        X.set_cell_vertical_center(c)
    left.text = ""
    add_image(left.paragraphs[0], ch["trend"], 2.3)
    prev = s["previous"]
    curr_score = s["overall_score"]
    delta = curr_score - prev["score"]
    delta_str = f"+{delta} 分" if delta >= 0 else f"{delta} 分"
    delta_color = "15803D" if delta >= 0 else "B91C1C"
    right.text = ""
    _r0 = right.paragraphs[0]
    add_run(_r0, "前次掃描　", size=10, color=T.GREY)
    add_run(_r0, f'{prev["date"]}　{prev["score"]} 分', size=10, color=T.SLATE, bold=True)
    p2 = right.add_paragraph()
    add_run(p2, "本次掃描　", size=10, color=T.GREY)
    add_run(p2, f'{s["scan_date"]}　{curr_score} 分', size=10, color=T.SLATE, bold=True)
    p3 = right.add_paragraph()
    add_run(p3, "變化　　　", size=10, color=T.GREY)
    add_run(p3, delta_str, size=11, color=delta_color, bold=True)
    if s.get("new_findings"):
        p4 = right.add_paragraph()
        add_run(p4, f'新出現 {len(s["new_findings"])} 項：{"、".join(s["new_findings"])}',
                size=9, color=T.GREY)
    _clear_table_borders(table)


def _clear_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


def _priorities(doc, data):
    h1(doc, "2", "優先處理清單")
    add_para(doc, [{"text": "以下依「嚴重度 × 影響範圍」排序，是本次最值得先處理的項目。編號對應第 4 章的完整說明，可直接跳轉查看「怎麼修」。",
                    "size": 10, "color": T.SLATE}], after=8, line=1.44)
    rows = []
    for i, pr in enumerate(data.get("priorities", []), 1):
        rows.append([str(i), ("chip", pr["severity"]), pr["problem"],
                     pr.get("category", ""), pr.get("ref", "")])
    if rows:
        data_table(doc, ["#", "嚴重度", "問題", "分類", "詳見"], rows,
                   [600, 1300, 3900, 1300, 1000])
    if data["summary"].get("priority_note"):
        add_para(doc, [{"text": data["summary"]["priority_note"], "size": 9, "color": T.GREY}],
                 before=6, after=0, line=1.38)


def _why_matters(doc, data):
    items = data.get("why_matters", [])
    if not items:
        return
    h1(doc, "3", "這些分類為什麼重要")
    add_para(doc, [{"text": "在逐項細節之前，先說明本次最弱的面向若不處理會有什麼實際後果，幫助你判斷投入的優先次序。",
                    "size": 10, "color": T.SLATE}], after=8, line=1.44)
    rows = [[it["category"], it["consequence"]] for it in items]
    data_table(doc, ["分類", "沒處理的話會怎樣"], rows, [2600, 5500])


def _findings(doc, data):
    h1(doc, "4", "發現項目")
    n = len(data["findings"])
    add_para(doc, [{"text": f"共 {n} 項，依嚴重度由高至低分組排列。每一張卡片獨立完整：問題是什麼、怎麼修、以及當下觀測到的檢測依據。各分類「為什麼重要」見第 3 章，修補後如何驗證見附錄 6.3。",
                    "size": 10, "color": T.SLATE}], after=6, line=1.44)
    by_sev = {}
    for f in data["findings"]:
        by_sev.setdefault(f["severity"], []).append(f)
    for sev in T.SEVERITY_ORDER:
        group = by_sev.get(sev, [])
        if not group:
            continue
        severity_group_header(doc, sev, len(group))
        for f in group:
            finding_card(doc, f)


def _scan_info(doc, data, ch):
    si = data.get("scan_info", {})
    h1(doc, "5", "掃描資訊與範圍")
    add_para(doc, [{"text": "本節說明這份報告涵蓋與未涵蓋的範圍，以及掃描當下擷取的網站畫面，供你確認判讀基礎。",
                    "size": 10, "color": T.SLATE}], after=8, line=1.44)
    h2(doc, "掃描範圍")
    scope = si.get("scope", {})
    rows = [[k, str(v)] for k, v in scope.items()]
    if rows:
        data_table(doc, ["項目", "設定"], rows, [2600, 5500])
    if si.get("warnings"):
        h2(doc, "掃描警示")
        for w in si["warnings"]:
            add_para(doc, [{"text": f"•  {w}", "size": 9.5, "color": T.SLATE}],
                     after=2, line=1.38)
    if si.get("screenshot"):
        h2(doc, "掃描當下的網站畫面")
        add_image(doc, si["screenshot"], 2.3)
        add_para(doc, [{"text": si.get("screenshot_caption", ""), "size": 8, "color": T.GREY}],
                 align=WD_ALIGN_PARAGRAPH.CENTER, after=0)


def _appendix(doc, data):
    ap = data.get("appendix", {})
    h1(doc, "6", "附錄")
    # glossary
    if ap.get("glossary"):
        h2(doc, "6.1　名詞解釋")
        add_para(doc, [{"text": "只列出本報告實際用到的專有名詞，方便對照閱讀。", "size": 9, "color": T.GREY}], after=5)
        rows = [[g["term"], g["explanation"]] for g in ap["glossary"]]
        data_table(doc, ["名詞", "白話說明"], rows, [1800, 6300])
    # tech index
    if ap.get("tech_index"):
        h2(doc, "6.2　技術索引")
        add_para(doc, [{"text": "供工程師與稽核人員對照。OWASP 為國際公認的網站安全風險分類，CWE 為通用軟體弱點編號；一般讀者可略過。",
                        "size": 9, "color": T.GREY}], after=5)
        rows = [[t["ref"], t["rule_id"], t.get("owasp_cwe", "— / —")] for t in ap["tech_index"]]
        data_table(doc, ["項次", "規則 ID", "OWASP / CWE"], rows, [1200, 4400, 2500])
    # verify
    h2(doc, "6.3　修補後如何驗證")
    add_para(doc, [{"text": ap.get("verify_note",
                    "完成修補後，重新執行一次 Argus 掃描，確認對應項目不再出現；下一份報告的摘要會列出這次解決了哪些項目。"),
                    "size": 10, "color": T.SLATE}], after=8, line=1.44)
    # authorization
    if ap.get("authorization"):
        h2(doc, "6.4　掃描授權聲明")
        auth = ap["authorization"]
        rows = [[k, str(v)] for k, v in auth.items()]
        data_table(doc, ["項目", "內容"], rows, [2600, 5500])
    # disclaimer
    h2(doc, "6.5　免責與報告產生方式")
    add_para(doc, [{"text": "本報告僅反映掃描當下、從網際網路可觀測到的外部特徵，不等同完整滲透測試或原始碼稽核，也不構成法律或合規意見。未列出的項目不代表不存在風險，實際修補請由具備權限的維運人員評估後執行。",
                    "size": 10, "color": T.SLATE}], after=6, line=1.44)
    add_para(doc, [{"text": "本報告採 Evidence-first 原則：SEO、AEO、GEO 與資安判斷全部來自爬蟲與規則引擎產生的可驗證證據，每一項發現都附上當下實際觀測到的內容，過程不使用 AI 改寫或推論結論。",
                    "size": 10, "color": T.SLATE}], after=0, line=1.44)
