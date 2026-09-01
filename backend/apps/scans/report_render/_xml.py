"""Low-level OOXML helpers python-docx doesn't expose directly."""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=130, right=130):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def set_cell_vertical_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), "center")
    tcPr.append(va)


def _border_el(edge, sz, color, space="0", val="single"):
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), val)
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    return e


def set_cell_borders(cell, sz=2, color="E2E8F0"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for edge in ("top", "start", "bottom", "end"):
        b.append(_border_el(edge, sz, color))
    tcPr.append(b)


def _ordered_insert(pPr, element, tag):
    """Insert element into pPr at the schema-correct position.
    pPr child order (subset): pStyle, keepNext, keepLines, pageBreakBefore,
    pBdr, shd, ..., spacing, ind, ... , jc, rPr, sectPr.
    We only manage a small set; place relative to known anchors."""
    # elements that must come AFTER our element (so we insert before the first present)
    after = {
        "w:pBdr": ["w:shd", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"],
        "w:shd":  ["w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr"],
        "w:ind":  ["w:jc", "w:rPr", "w:sectPr"],
    }.get(tag, [])
    for a in after:
        anchor = pPr.find(qn(a))
        if anchor is not None:
            anchor.addprevious(element)
            return
    pPr.append(element)


def set_para_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    _ordered_insert(pPr, shd, "w:shd")


def set_para_borders(paragraph, edges):
    """edges: dict edge-> (sz,color,space). Order enforced: top,left,bottom,right."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        if edge in edges:
            sz, color, space = edges[edge]
            pBdr.append(_border_el(edge, sz, color, space))
    _ordered_insert(pPr, pBdr, "w:pBdr")


def set_para_indent(paragraph, left=None, right=None):
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        _ordered_insert(pPr, ind, "w:ind")
    if left is not None:
        ind.set(qn("w:left"), str(left))
    if right is not None:
        ind.set(qn("w:right"), str(right))


def set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:cantSplit")
    trPr.append(e)


def add_page_number_field(paragraph):
    """Insert a PAGE field into a paragraph run."""
    run = paragraph.add_run()
    fldSimple = OxmlElement("w:fldSimple")
    fldSimple.set(qn("w:instr"), "PAGE")
    run._r.addprevious(fldSimple)


def add_watermark_to_header(header, text="ARGUS", color="0C4A6E", opacity=".06"):
    """Inject a diagonal VML WordArt watermark that repeats on every page."""
    p = header.add_paragraph()
    r = p.add_run()
    pict_xml = f'''<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
      xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800"
        path="m@7,l@8,m@5,21600l@6,21600e">
        <v:formulas><v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>
        <v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/>
        <v:f eqn="if @0 @3 0"/><v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/>
        <v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/>
        <v:f eqn="mid @7 @8"/><v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/></v:formulas>
        <v:path textpathok="t" o:connecttype="custom"
          o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
        <v:textpath on="t" fitshape="t"/></v:shapetype>
      <v:shape id="ArgusWM" o:spid="_x0000_s2049" type="#_x0000_t136"
        style="position:absolute;margin-left:0;margin-top:0;width:520pt;height:170pt;rotation:315;
        z-index:-251654144;mso-position-horizontal:center;mso-position-horizontal-relative:margin;
        mso-position-vertical:center;mso-position-vertical-relative:margin"
        o:allowincell="f" fillcolor="#{color}" stroked="f">
        <v:fill opacity="{opacity}"/>
        <v:textpath style="font-family:&quot;Microsoft JhengHei&quot;;font-weight:bold;v-text-align:center"
          string="{text}"/></v:shape></w:pict>'''
    from docx.oxml import parse_xml
    r._r.append(parse_xml(pict_xml))
    return p


def set_page_size_letter(section):
    section.page_width = 12240 * 635   # will be overridden; kept for clarity


def fix_zoom_setting(doc):
    """python-docx emits <w:zoom/> without the required percent attr; set it."""
    from docx.oxml.ns import qn as _qn
    settings = doc.settings.element
    zoom = settings.find(_qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        settings.insert(0, zoom)
    zoom.set(_qn("w:percent"), "100")
