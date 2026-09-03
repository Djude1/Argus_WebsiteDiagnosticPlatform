"""報告可讀性與排版測試（稽核第三階段：C1-C5、D1-D5）。

背景（docs/scan-report-quality-audit-2026-08-30.md）：
scan 25 的 .docx 是 328 段純文字流，其中 293 段是 Normal 樣式——零表格、零分頁、
無封面、無目錄、無頁碼、無顏色，且直接把 rule_id、evidence_source=rule_engine
這類內部識別碼印給不懂技術的中小企業主看。
"""

from __future__ import annotations

from django.contrib.auth import get_user_model
from django.test import TestCase
from docx import Document

from apps.scans.models import Finding, Page, ScanJob
from apps.scans.reports import build_scan_report

User = get_user_model()


class ReportLayoutTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="layout", password="safe-test-password")
        self.scan_job = ScanJob.objects.create(
            user=self.user,
            original_url="https://example.com/",
            normalized_url="https://example.com/",
            origin="https://example.com",
            status=ScanJob.Status.COMPLETED,
            max_depth=2, max_pages=20,
            overall_score=61,
            category_scores={"seo": 70, "aeo": 100, "geo": 59, "security": 39},
            top_actions=[{"title": "缺少 HSTS", "category": "security", "severity": "high"}],
        )
        page = Page.objects.create(
            scan_job=self.scan_job, url="https://example.com/a",
            final_url="https://example.com/a", origin="https://example.com",
        )
        Finding.objects.create(
            scan_job=self.scan_job, page=page, severity="high",
            category=Finding.Category.SECURITY, title="缺少 HSTS",
            description="回應未帶 Strict-Transport-Security。",
            remediation="於 web server 加上 HSTS 標頭。",
            evidence="回應標頭查無 strict-transport-security",
            rule_id="header-hsts-missing", owasp_category="A05", cwe_id="CWE-319",
            evidence_source="rule_engine", evidence_type="text",
            ai_handoff_prompt="我網站有以下問題…", priority_score=75.0,
        )

    def _doc(self) -> Document:
        return Document(build_scan_report(self.scan_job))

    def _text(self) -> str:
        doc = self._doc()
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    # --- D 排版 ------------------------------------------------------
    def test_report_uses_tables(self):
        """全文用段落堆疊，資訊沒有結構可掃讀。"""
        self.assertGreater(len(self._doc().tables), 0)

    def test_report_is_visually_structured_with_charts(self):
        """報告要有視覺結構，不能是一條純文字流。

        舊版用分頁數當代理指標（因為當時只有段落可用）。report_render 改用
        程式繪製的圖表 + 卡片式版面，分頁由內容長度自然決定，數分頁已無意義；
        改鎖「有圖表」——那才是這次採用新排版真正要拿到的東西。
        """
        doc = self._doc()

        # 分數環圈、分類長條、嚴重度分佈至少三張
        self.assertGreaterEqual(len(doc.inline_shapes), 3)
        self.assertGreater(len(doc.tables), 0)

    def test_report_has_a_branded_cover_and_a_summary_first_page(self):
        """封面要有品牌，第一章要能 30 秒看懂整體狀況。

        2026-09-01 起排版交給 report_render。舊版的「目錄」章節被「一頁摘要」
        取代——後者對網站主更有用（目錄只是導航，摘要直接回答「我的網站現在
        怎麼樣」），且與使用者提供的範本一致。
        """
        text = self._text()

        self.assertIn("ARGUS", text)
        self.assertIn("一頁摘要", text)
        self.assertIn("整體分數", text)

    def test_report_has_header_and_footer_with_page_number(self):
        section = self._doc().sections[0]
        header_text = "\n".join(p.text for p in section.header.paragraphs)
        footer_xml = "".join(p._p.xml for p in section.footer.paragraphs)
        self.assertIn("ARGUS", header_text)
        self.assertIn("PAGE", footer_xml)   # 頁碼 field code

    def test_severity_is_colour_coded(self):
        """嚴重度只有文字，掃讀時分不出輕重。"""
        doc = self._doc()
        coloured = [
            run for p in doc.paragraphs for run in p.runs
            if run.font.color is not None and run.font.color.rgb is not None
        ]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    coloured += [
                        r for p in cell.paragraphs for r in p.runs
                        if r.font.color is not None and r.font.color.rgb is not None
                    ]
        self.assertGreater(len(coloured), 0)

    def test_report_shows_which_categories_the_problems_cluster_in(self):
        """報告要能回答「問題集中在哪一類」，不只是「哪一類做得多好」。

        前端一直有這張佔比圖，報告沒有。兩張圖互補：分類分數看體質，
        佔比看問題分佈。
        """
        for index in range(3):
            Finding.objects.create(
                scan_job=self.scan_job, page=None, severity="low",
                category=Finding.Category.SEO, title=f"SEO 問題 {index}",
                description="d", remediation="r", rule_id=f"seo-{index}",
                ai_handoff_prompt="p", priority_score=25.0,
            )

        text = self._text()

        self.assertIn("問題集中在哪些分類", text)
        self.assertIn("SEO 搜尋引擎最佳化 3 項", text)

    def test_category_share_uses_the_same_colours_as_the_dashboard(self):
        """同一份掃描在畫面與報告上，分類顏色必須一致。"""
        from apps.scans.report_render import theme

        self.assertEqual(theme.category_color("SEO 搜尋引擎最佳化"), "6366F1")
        self.assertEqual(theme.category_color("資訊安全"), "EF4444")
        # 未知分類不得炸掉，退回中性灰
        self.assertEqual(theme.category_color("未來的新分類"), theme.GREY)

    # --- C 可讀性與術語 -----------------------------------------------
    def test_internal_identifiers_are_not_shown_in_the_findings_body(self):
        """rule_engine / text 這種內部欄位值對使用者零意義。"""
        text = self._text()
        self.assertNotIn("rule_engine", text)
        self.assertNotIn("證據型態", text)

    def test_rule_id_moves_to_a_technical_index(self):
        """rule_id 仍要可查（技術人員需要），但不該混在正文裡。"""
        doc = self._doc()
        body = "\n".join(p.text for p in doc.paragraphs)
        table_text = "\n".join(
            cell.text for t in doc.tables for row in t.rows for cell in row.cells
        )
        self.assertNotIn("header-hsts-missing", body)
        self.assertIn("header-hsts-missing", table_text)

    def test_english_jargon_heading_is_translated(self):
        self.assertNotIn("Deterministic Evidence", self._text())
        self.assertIn("檢測依據", self._text())

    def test_report_answers_what_why_how_and_verification(self):
        """報告必須回答這四個問題——但只講一次。

        原本的要求是「每一項發現都要有這四段」。scan 38 的實測顯示那造成大量重複
        （「為什麼要在意」出現 14 次卻只有 6 種內容），34 頁裡有一大半是樣板。
        所以要求改成：屬於分類的講一次（摘要）、屬於流程的講一次（附錄）、
        屬於這一項的才逐項列（見 tests_report_compactness.py）。
        """
        text = self._text()

        self.assertIn("問題是什麼", text)       # 逐項
        self.assertIn("怎麼修", text)           # 逐項
        self.assertIn("這些分類為什麼重要", text)  # 摘要，講一次
        self.assertIn("修補後如何驗證", text)      # 附錄，講一次

    def test_positive_info_finding_is_not_described_as_a_threat(self):
        """info 常是正向指標，不能套用「會被攻擊者利用」那套說法。

        實際產出報告時發現：「Nuclei 探針受 WAF 攔截」（代表防護有效）底下寫著
        「這類問題會被攻擊者利用，可能導致網站被入侵」——與該項的意義完全相反。
        """
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="info",
            category=Finding.Category.SECURITY, title="Nuclei 探針受 WAF 攔截",
            description="偵測到 Cloudflare 保護，屬正向安全指標。",
            remediation="無需修復。", rule_id="waf-detected",
            ai_handoff_prompt="p", priority_score=10.0,
        )
        self.scan_job.findings.exclude(rule_id="waf-detected").delete()

        text = self._text()

        self.assertNotIn("會被攻擊者利用", text)
        self.assertNotIn("導致網站被入侵", text)

    def test_info_finding_does_not_promise_re_scan_verification(self):
        """info 項目不會在下次掃描「消失」，不該叫讀者去確認它不見了。"""
        self.scan_job.findings.all().delete()
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="info",
            category=Finding.Category.SECURITY, title="Nuclei 探針受 WAF 攔截",
            description="偵測到 Cloudflare 保護。", remediation="無需修復。",
            rule_id="waf-detected", ai_handoff_prompt="p", priority_score=10.0,
        )

        self.assertNotIn("修好了怎麼確認", self._text())

    def test_actionable_info_finding_still_shows_how_to_fix(self):
        """info 不全是正向指標。

        scan 28 的 5 個 info 項目裡只有 1 個是正向（WAF 攔截），其餘 4 個
        （缺 X-Content-Type-Options、缺 canonical…）都是可以改的小問題。
        對它們宣告「不需要採取任何修補動作」再印出修補方式，是自相矛盾。
        """
        self.scan_job.findings.all().delete()
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="info",
            category=Finding.Category.SECURITY, title="缺少 X-Content-Type-Options",
            description="回應標頭缺少 x-content-type-options。",
            remediation="設定 x-content-type-options: nosniff。",
            rule_id="header-xcto", ai_handoff_prompt="p", priority_score=10.0,
        )

        text = self._text()

        self.assertNotIn("不需要採取任何修補動作", text)
        self.assertIn("怎麼修", text)
        self.assertIn("設定 x-content-type-options: nosniff。", text)

    def test_report_drops_description_warnings_that_contradict_its_masking(self):
        """scanner 的 description 帶著「此項目顯示原始個資」，但報告會遮罩。

        那句對前端（依使用者要求顯示原始個資）成立，搬進有遮罩的報告就是假話；
        報告本來就會在「檢測依據」下輸出自己那句正確的遮罩提示。
        """
        self.scan_job.findings.all().delete()
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="high",
            category=Finding.Category.SECURITY, title="頁面外洩個人資料 (PII)",
            description=(
                "⚠️ 此項目顯示原始個資，請依個資法妥善處理本報告。\n"
                "在頁面內容中偵測到 1 筆疑似個資。"
            ),
            remediation="下架該頁。", evidence="台灣手機：0912345678",
            rule_id="pii", ai_handoff_prompt="p", priority_score=75.0,
        )

        text = self._text()

        self.assertNotIn("此項目顯示原始個資", text)
        self.assertIn("在頁面內容中偵測到 1 筆疑似個資。", text)

    def test_security_verification_advice_is_not_header_specific(self):
        """並非所有資安問題都靠檢查回應標頭驗證（例如頁面外洩個資）。"""
        self.scan_job.findings.all().delete()
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="high",
            category=Finding.Category.SECURITY, title="頁面外洩個人資料",
            description="偵測到個資。", remediation="下架該頁。",
            rule_id="pii", ai_handoff_prompt="p", priority_score=75.0,
        )

        self.assertNotIn("curl -I", self._text())

    def test_glossary_explains_terms_that_appear_in_this_report(self):
        text = self._text()
        self.assertIn("名詞解釋", text)
        self.assertIn("HSTS", text)
        self.assertIn("強制", text)  # HSTS 的中文解釋

    def test_glossary_omits_terms_not_present_in_this_report(self):
        """詞彙表要對應這份報告的內容，不是貼一份固定清單。"""
        self.assertNotIn("DNSSEC", self._text())
