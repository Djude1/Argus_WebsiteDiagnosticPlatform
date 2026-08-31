"""報告精簡度（使用者回報：34 頁太長、樣板重複讀起來像灌水）。

scan 38 的實測分布：發現項目佔全文 78.9%，其中
  - 「想更深入了解？」的 AI 提示詞 5412 字（佔發現項目 43%），內容是把上方
    「問題是什麼／檢測依據／怎麼修」原封不動再抄一遍
  - 「為什麼要在意」出現 14 次但只有 6 種內容
  - 「修好了怎麼確認」出現 14 次但只有 3 種內容

只跟分類有關的樣板要講一次，每項發現只留屬於它自己的內容。
"""

from __future__ import annotations

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.utils import timezone
from docx import Document

from apps.scans.models import Finding, ScanJob
from apps.scans.reports import build_scan_report

User = get_user_model()


class ReportCompactnessTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="compact", password="safe-test-password")
        self.scan_job = ScanJob.objects.create(
            user=self.user, original_url="https://example.com/",
            normalized_url="https://example.com/", origin="example.com",
            status=ScanJob.Status.COMPLETED, overall_score=70,
            category_scores={"security": 40, "seo": 80},
            completed_at=timezone.now(),
        )
        for index in range(6):
            Finding.objects.create(
                scan_job=self.scan_job, page=None, severity="medium",
                category=Finding.Category.SECURITY, title=f"資安問題 {index}",
                description=f"描述 {index}", remediation=f"修補 {index}",
                evidence=f"證據 {index}", rule_id=f"sec-{index}",
                ai_handoff_prompt=f"我網站有以下問題…{index}", priority_score=50.0,
            )

    def _doc(self):
        return Document(build_scan_report(self.scan_job))

    def _text(self) -> str:
        document = self._doc()
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def test_ai_prompt_is_not_repeated_under_every_finding(self):
        """提示詞內容與上方三段完全重疊，逐項再印一次是純粹的重複。"""
        text = self._text()

        self.assertNotIn("我網站有以下問題…0", text)
        self.assertNotIn("想更深入了解？", text)

    def test_report_still_tells_the_reader_how_to_use_ai(self):
        """功能不能因為精簡就消失——改成在附錄說明一次怎麼做。"""
        text = self._text()

        self.assertIn("ChatGPT", text)
        self.assertEqual(text.count("ChatGPT"), 1)

    def test_category_impact_is_stated_once_not_per_finding(self):
        text = self._text()

        self.assertIn("會被攻擊者利用", text)
        self.assertEqual(text.count("會被攻擊者利用"), 1)

    def test_verification_advice_is_stated_once_not_per_finding(self):
        text = self._text()

        self.assertEqual(text.count("重新執行一次 Argus 掃描"), 1)

    def test_each_finding_keeps_its_own_specific_content(self):
        """精簡不能砍掉每項發現真正屬於自己的資訊。"""
        text = self._text()

        for index in range(6):
            self.assertIn(f"描述 {index}", text)
            self.assertIn(f"修補 {index}", text)
            self.assertIn(f"證據 {index}", text)

    def test_findings_do_not_each_carry_their_own_table(self):
        """每項發現一張中繼表格，在 Word 裡非常吃垂直空間。"""
        document = self._doc()

        # 摘要 3 張 + 目錄 + 掃描範圍 + 優先建議 + 附錄數張，但不該隨 finding 數成長
        self.assertLess(len(document.tables), 6 + 6)

    def test_affected_pages_list_is_capped(self):
        from apps.scans.models import Page

        pages = [
            Page.objects.create(
                scan_job=self.scan_job, url=f"https://example.com/p{i}",
                final_url=f"https://example.com/p{i}", origin="example.com", depth=i,
            )
            for i in range(12)
        ]
        for page in pages:
            Finding.objects.create(
                scan_job=self.scan_job, page=page, severity="low",
                category=Finding.Category.SEO, title="標題過長",
                description="d", remediation="r", rule_id="seo-title",
                ai_handoff_prompt="p", priority_score=25.0,
            )

        findings_section = self._text().split("5\u3000附錄")[0]

        listed = findings_section.count("https://example.com/p")
        self.assertEqual(listed, 5)         # 只列前 5 個
        self.assertIn("另 7 處", findings_section)   # 其餘收成一句話
