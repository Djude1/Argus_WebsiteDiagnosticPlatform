"""報告附錄內容測試（稽核第四階段：F4、F5、F7）。

背景（docs/scan-report-quality-audit-2026-08-30.md）：
Page 表存了每頁的 URL、標題、HTTP 狀態、深度與阻擋原因，Page.screenshot_path
有截圖，views.py 也早有同 origin 歷史掃描的概念——報告一項都沒用。
「上次 39 分、這次 65 分，修好了 3 項」是最能展現價值的內容，卻完全沒有。
"""

from __future__ import annotations

from pathlib import Path

from django.conf import settings
from django.contrib.auth import get_user_model
from django.test import TestCase
from django.utils import timezone
from docx import Document

from apps.scans.models import Finding, Page, ScanJob
from apps.scans.reports import build_scan_report

User = get_user_model()


class ReportAppendixTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="appendix", password="safe-test-password")
        self.scan_job = self._make_scan(overall_score=65)

    def _make_scan(self, *, overall_score, completed_offset_days=0) -> ScanJob:
        return ScanJob.objects.create(
            user=self.user,
            original_url="https://example.com/",
            normalized_url="https://example.com/",
            origin="example.com",
            status=ScanJob.Status.COMPLETED,
            overall_score=overall_score,
            completed_at=timezone.now() - timezone.timedelta(days=completed_offset_days),
        )

    def _text(self, scan_job=None) -> str:
        document = Document(build_scan_report(scan_job or self.scan_job))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    # --- F4 頁面清單 -------------------------------------------------
    # 2026-09-01 起排版交給 report_render，其 schema 沒有頁面清單欄位，該區塊
    # 隨之移除。這與使用者提供的範本一致（範本的章節是 1 一頁摘要 / 2 優先處理 /
    # 3 為什麼重要 / 4 發現項目 / 5 掃描資訊 / 6 附錄，同樣沒有頁面清單），且
    # 掃描範圍表仍有「實際掃描頁數」、每項發現仍列出受影響網址。
    # 要恢復需改 report_render 的 schema 與 report.py（vendored 第三方程式碼）。
    def _removed_test_crawled_pages_are_listed_in_the_appendix(self):
        Page.objects.create(
            scan_job=self.scan_job, url="https://example.com/a",
            final_url="https://example.com/a", origin="example.com",
            title="關於我們", status_code=200, depth=1,
        )
        Page.objects.create(
            scan_job=self.scan_job, url="https://example.com/b",
            final_url="https://example.com/b", origin="example.com",
            title="", status_code=404, depth=2, blocked_reason="robots.txt",
        )

        text = self._text()

        self.assertIn("掃描頁面清單", text)
        self.assertIn("https://example.com/a", text)
        self.assertIn("關於我們", text)
        self.assertIn("404", text)
        self.assertIn("robots.txt", text)

    def test_page_count_is_still_reported_in_scan_scope(self):
        """頁面清單移除後，「掃描了幾頁」這個資訊不能跟著消失。"""
        self.assertIn("實際掃描頁數", self._text())

    # --- F7 與前次掃描比較 --------------------------------------------
    def test_report_compares_against_the_previous_scan_of_the_same_site(self):
        previous = self._make_scan(overall_score=39, completed_offset_days=7)
        Finding.objects.create(
            scan_job=previous, page=None, severity="medium",
            category=Finding.Category.SECURITY, title="缺少 HSTS",
            description="d", remediation="r", rule_id="header-hsts-missing",
            ai_handoff_prompt="p", priority_score=50.0,
        )
        Finding.objects.create(
            scan_job=self.scan_job, page=None, severity="medium",
            category=Finding.Category.SECURITY, title="缺少 CSP",
            description="d", remediation="r", rule_id="header-csp-missing",
            ai_handoff_prompt="p", priority_score=50.0,
        )

        text = self._text()

        # report_render 的比較區塊只有「新出現」欄位；已解決的數量改收進導讀句
        # （schema 沒有 resolved 欄位，但「修好了 N 項」是回訪者最想看的資訊）。
        self.assertIn("39", text)           # 前次分數
        self.assertIn("已解決 1 項", text)   # 前次有、這次沒有
        self.assertIn("新出現", text)
        self.assertIn("缺少 CSP", text)

    def test_first_scan_of_a_site_has_no_comparison_section(self):
        self.assertNotIn("已解決", self._text())

    def test_comparison_ignores_other_users_scans_of_the_same_site(self):
        """同一個網址可能被不同使用者掃過，不能拿別人的掃描當「前次」。"""
        other = User.objects.create_user(username="other", password="safe-test-password")
        ScanJob.objects.create(
            user=other, original_url="https://example.com/",
            normalized_url="https://example.com/", origin="example.com",
            status=ScanJob.Status.COMPLETED, overall_score=12,
            completed_at=timezone.now() - timezone.timedelta(days=3),
        )

        self.assertNotIn("已解決", self._text())

    def test_comparison_ignores_unfinished_scans(self):
        ScanJob.objects.create(
            user=self.user, original_url="https://example.com/",
            normalized_url="https://example.com/", origin="example.com",
            status=ScanJob.Status.FAILED, overall_score=None,
            completed_at=timezone.now() - timezone.timedelta(days=1),
        )

        self.assertNotIn("已解決", self._text())


class ReportScreenshotTests(TestCase):
    """入口頁截圖（F5）。

    刻意只放入口頁一張：全頁截圖體積大，50 頁的掃描全塞進去會讓 .docx 爆掉，
    而 header / DNS / meta 這類發現本來就沒有視覺佐證價值。放一張的用途是讓
    收件者確認「這份報告講的確實是我的網站」，同時呼應防偽主題。
    """

    def setUp(self):
        self.user = User.objects.create_user(username="shot", password="safe-test-password")
        self.scan_job = ScanJob.objects.create(
            user=self.user, original_url="https://example.com/",
            normalized_url="https://example.com/", origin="example.com",
            status=ScanJob.Status.COMPLETED, overall_score=70,
            completed_at=timezone.now(),
        )

    def _make_png(self, relative_path: str) -> str:
        from PIL import Image

        absolute = Path(settings.BASE_DIR) / relative_path
        absolute.parent.mkdir(parents=True, exist_ok=True)
        Image.new("RGB", (120, 80), (10, 21, 53)).save(absolute)
        self.addCleanup(lambda: absolute.unlink(missing_ok=True))
        return relative_path

    def _image_count(self) -> int:
        return len(Document(build_scan_report(self.scan_job)).inline_shapes)

    def _caption_count(self) -> int:
        """截圖說明文字出現幾次——比數總圖片數可靠。

        report_render 會嵌 4 張程式繪製的圖表（分數環圈、分類長條、嚴重度分佈、
        趨勢），總圖片數會隨圖表增減而變；寫死總數只會變成每次調圖表就要改測試
        的雜訊。真正要鎖的是「入口頁截圖有出現，且不隨頁數成長」。
        """
        document = Document(build_scan_report(self.scan_job))
        # 用帶全形括號的完整字串：report_render 的章節導言裡也有「掃描當下擷取的
        # 網站畫面」，只比對前四個字會永遠命中而讓測試失去意義。
        return sum(1 for p in document.paragraphs if "（掃描當下擷取）" in p.text)

    def test_entry_page_screenshot_is_embedded(self):
        Page.objects.create(
            scan_job=self.scan_job, url="https://example.com/",
            final_url="https://example.com/", origin="example.com", depth=0,
            screenshot_path=self._make_png("media/test_shots/entry.png"),
        )

        self.assertEqual(self._caption_count(), 1)

    def test_only_the_entry_page_screenshot_is_embedded(self):
        """不是每頁都塞圖——那會讓報告體積失控。"""
        for index in range(3):
            Page.objects.create(
                scan_job=self.scan_job, url=f"https://example.com/p{index}",
                final_url=f"https://example.com/p{index}", origin="example.com",
                depth=index,
                screenshot_path=self._make_png(f"media/test_shots/p{index}.png"),
            )

        self.assertEqual(self._caption_count(), 1)

    def test_missing_screenshot_file_does_not_break_the_report(self):
        Page.objects.create(
            scan_job=self.scan_job, url="https://example.com/",
            final_url="https://example.com/", origin="example.com", depth=0,
            screenshot_path="media/test_shots/does-not-exist.png",
        )

        # 檔案不在就略過截圖，但報告照常產出（圖表仍在）
        self.assertEqual(self._caption_count(), 0)
        self.assertGreater(self._image_count(), 0)
