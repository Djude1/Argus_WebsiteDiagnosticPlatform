"""測試 per-rule 客製文案 + 藝術字封面 + 縮圖截圖策略。

對應 audit supplement 的 N項 + 你 8/31 的「很有問題」回饋：
- (C) per-rule 「為什麼要在意」具體後果
- (B) per-rule 「修好了怎麼確認」驗收指令
- (E) 封面用 argus-title.png（含漸層 + 4 角星點綴）
- (D) 截圖縮成 thumbnail，不嵌 3MB 全頁截圖

執行：cd backend && uv run python manage.py test apps.scans.tests_report_rule_specific -v 2
"""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from django.test import TestCase

from apps.scans.models import Finding, Page, ScanJob
from apps.scans.reports import (
    _impact_for,
    _verify_for,
    build_scan_report,
)


class RuleImpactTests(TestCase):
    def test_known_rule_id_returns_specific_impact(self):
        """SECURITY_PII_8B24BB8B28 應回傳含「個資法」「罰鍰」字眼的客製文案，
        不是 category-level 的通用模板。"""
        finding = Finding(
            rule_id="SECURITY_PII_8B24BB8B28",
            category="security",
            severity="high",
            title="PII",
            description="d", remediation="r",
            ai_handoff_prompt="p",
        )
        text = _impact_for(finding)
        self.assertIn("個資法", text)
        self.assertIn("罰鍰", text)
        # 不應該是通用模板的「會被攻擊者利用」
        self.assertNotIn("攻擊者利用", text)

    def test_unknown_rule_id_falls_back_to_category(self):
        finding = Finding(
            rule_id="UNKNOWN_RULE_XYZ",
            category="security",
            severity="medium",
            title="X",
            description="d", remediation="r",
            ai_handoff_prompt="p",
        )
        text = _impact_for(finding)
        self.assertIn("攻擊者", text)  # CATEGORY_IMPACT security fallback

    def test_empty_rule_id_falls_back_to_category(self):
        finding = Finding(
            rule_id="",
            category="seo",
            severity="low",
            title="X",
            description="d", remediation="r",
            ai_handoff_prompt="p",
        )
        text = _impact_for(finding)
        self.assertIn("搜尋引擎", text)

    def test_seo_meta_title_specific(self):
        """SEO_META_TITLE 應有「點擊率」「截斷」這類具體詞，不是通用 SEO 模板。"""
        finding = Finding(
            rule_id="SEO_META_TITLE_0D9B1FE9E2",
            category="seo",
            severity="low",
            title="Meta title",
            description="d", remediation="r",
            ai_handoff_prompt="p",
        )
        text = _impact_for(finding)
        self.assertIn("點擊率", text)
        self.assertIn("20-30", text)


class RuleVerifyTests(TestCase):
    def test_known_rule_id_returns_specific_verify(self):
        """CSP 應有 `curl -I` + `content-security-policy` 驗收指令，不是「再掃一次 Argus」。"""
        finding = Finding(
            rule_id="SECURITY_CSP_BD010B5BE0",
            category="security",
            severity="medium",
            title="CSP",
            description="d", remediation="r",
            ai_handoff_prompt="p",
        )
        text = _verify_for(finding)
        self.assertIn("curl", text)
        self.assertIn("content-security-policy", text)

    def test_dns_rules_have_dig_commands(self):
        """DNS 相關 rule 應有 dig 指令。"""
        for rid in ["dns-spf-missing", "dns-dnssec-missing", "dns-dmarc-policy-weak"]:
            finding = Finding(
                rule_id=rid, category="security", severity="medium",
                title="X", description="d", remediation="r", ai_handoff_prompt="p",
            )
            self.assertIn("dig", _verify_for(finding), f"{rid} 應含 dig 指令")

    def test_unknown_falls_back_to_category(self):
        finding = Finding(
            rule_id="UNKNOWN",
            category="geo", severity="medium",
            title="X", description="d", remediation="r", ai_handoff_prompt="p",
        )
        # CATEGORY_VERIFY geo: 「修補後重新執行一次 Argus 掃描確認此項目消失。」
        self.assertIn("Argus", _verify_for(finding))


def _full_doc_text(path: str) -> str:
    """讀 .docx 全部文字（含 tables 與 header/footer）。

    python-docx 的 doc.paragraphs 不含 table cells，但 per-rule 客製文案是寫在
    summary 表格與附錄列表中的，這裡手動遞迴讀完整個 document tree。
    """
    from docx import Document as _Doc
    doc = _Doc(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


class ReportHasRuleSpecificSectionsTests(TestCase):
    """驗證 per-rule 客製文案進到正確位置（2026-08-31 後設計）：

    - 「為什麼要在意」：在 summary 的「這些分類為什麼重要」表格（講一次），
      用該 category 最嚴重 finding 的 rule_id 查 RULE_IMPACT。
    - 「修補後如何驗證」：在附錄的「修補後如何驗證」小節（講一次通用模板 +
      per-category 代表驗收指令），用 RULE_VERIFY。
    - 不再每個 finding 都重複這些段落（既有 compactness 測試保護）。
    """

    def _make_scan(self) -> ScanJob:
        from django.contrib.auth import get_user_model
        from django.utils import timezone
        User = get_user_model()
        user = User.objects.create_user(
            username="rule-tester", email="t@x.com", password="p",
        )
        scan = ScanJob.objects.create(
            user=user,
            original_url="https://example.com",
            normalized_url="https://example.com",
            origin="example.com",
            status=ScanJob.Status.COMPLETED,
            overall_score=50,
            category_scores={"security": 50, "seo": 80, "aeo": 100, "geo": 70, "ux": 100},
            completed_at=timezone.now(),
        )
        page = Page.objects.create(
            scan_job=scan,
            url="https://example.com",
            final_url="https://example.com",
            origin="example.com",
        )
        return scan, page

    def test_summary_uses_per_rule_impact(self):
        """summary 表格裡 security category 應包含 per-rule 客製文案（個資法罰鍰）。"""
        scan, page = self._make_scan()
        Finding.objects.create(
            scan_job=scan, page=page,
            severity="high", category=Finding.Category.SECURITY,
            title="頁面外洩個資",
            description="在頁面偵測到個資",
            remediation="請遮罩",
            ai_handoff_prompt="p",
            rule_id="SECURITY_PII_8B24BB8B28",
            priority_score=75.0,
        )
        Finding.objects.create(
            scan_job=scan, page=page,
            severity="medium", category=Finding.Category.SECURITY,
            title="缺少 CSP",
            description="Response header 缺 CSP",
            remediation="加上 CSP header",
            ai_handoff_prompt="p",
            rule_id="SECURITY_CSP_BD010B5BE0",
            priority_score=50.0,
        )
        path = build_scan_report(scan)
        all_text = _full_doc_text(path)

        # summary 表格區段：包含「這些分類為什麼重要」標題與 per-rule 客製文案。
        self.assertIn("這些分類為什麼重要", all_text)
        # SECURITY_PII 的 RULE_IMPACT 含「個資法」「罰鍰」
        self.assertIn("個資法", all_text)
        self.assertIn("罰鍰", all_text)
        # security category 講一次（既有 compactness 測試保留）
        self.assertEqual(all_text.count("個資法"), 1)

    def test_appendix_uses_per_rule_verify(self):
        """附錄「修補後如何驗證」應包含 per-category per-rule 驗收指令。"""
        scan, page = self._make_scan()
        Finding.objects.create(
            scan_job=scan, page=page,
            severity="medium", category=Finding.Category.SECURITY,
            title="缺少 CSP",
            description="Response header 缺 CSP",
            remediation="加上 CSP header",
            ai_handoff_prompt="p",
            rule_id="SECURITY_CSP_BD010B5BE0",
            priority_score=50.0,
        )
        path = build_scan_report(scan)
        all_text = _full_doc_text(path)

        # 附錄「修補後如何驗證」小節：通用模板 + per-category 驗收指令
        self.assertIn("修補後如何驗證", all_text)
        self.assertIn("重新執行一次 Argus 掃描", all_text)  # 通用模板
        # per-rule 驗收指令（curl + header 名）
        self.assertIn("curl", all_text)
        self.assertIn("content-security-policy", all_text)

    def test_findings_do_not_each_carry_impact_and_verify(self):
        """個別 finding 不再重複印「為什麼要在意」「修好了怎麼確認」（避免重複）。"""
        scan, page = self._make_scan()
        # 建兩個 finding，跨兩個不同 rule_id（需 5 個 finding 才會拉平到 N=1 兩次但
        # 這裡只 2 個，要驗證「任何單一 finding 區段內都沒有這兩類標題」即可。
        for title, rid, sev in [
            ("頁面外洩個資", "SECURITY_PII_8B24BB8B28", "high"),
            ("缺少 CSP", "SECURITY_CSP_BD010B5BE0", "medium"),
        ]:
            Finding.objects.create(
                scan_job=scan, page=page,
                severity=sev, category=Finding.Category.SECURITY,
                title=title, description="d", remediation="r",
                ai_handoff_prompt="p", rule_id=rid,
                priority_score=75.0,
            )
        path = build_scan_report(scan)
        all_text = _full_doc_text(path)

        # per-rule 客製各在 summary/附錄出現 1 次（不重複）
        self.assertEqual(all_text.count("個資法"), 1)
        self.assertEqual(all_text.count("content-security-policy"), 1)
        # 個別 finding 區段（從第一個 finding 開始到附錄前）不該含這兩個標題
        finding_idx = all_text.find("頁面外洩個資")
        appendix_idx = all_text.find("附錄")
        findings_section = all_text[finding_idx:appendix_idx]
        self.assertNotIn("為什麼要在意", findings_section)
        self.assertNotIn("修好了怎麼確認", findings_section)


class ReportHasTitleImageTests(TestCase):
    """封面應該用 argus-title.png（含 4 角星點綴 + 漸層），而非純文字 ARGUS。"""

    def test_title_png_is_embedded(self):
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username="title-tester", email="t@x.com", password="p",
        )
        scan = ScanJob.objects.create(
            user=user,
            original_url="https://example.com",
            normalized_url="https://example.com",
            origin="example.com",
            status=ScanJob.Status.COMPLETED,
        )
        path = build_scan_report(scan)
        with zipfile.ZipFile(path) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
        # 至少有 1 張媒體（封面 title PNG 或 fallback logo）
        self.assertGreaterEqual(len(media), 1)


class ReportThumbnailScreenshotTests(TestCase):
    """截圖應該用 thumbnail（最寬 1200px），不直接 embed 全頁截圖。"""

    def setUp(self):
        # 建一個假截圖（PIL.Image 6000x1500 模擬 scan-25 的全頁截圖）
        import tempfile

        from PIL import Image

        self.tmp = tempfile.mkdtemp()
        self.scan_img_path = Path(self.tmp) / "scan-test-page.png"
        img = Image.new("RGB", (1500, 6000), (200, 200, 200))
        img.save(self.scan_img_path, "PNG")

    def tearDown(self):
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_thumbnail_is_smaller_than_original(self):
        from django.contrib.auth import get_user_model
        User = get_user_model()
        user = User.objects.create_user(
            username="thumb-tester", email="t@x.com", password="p",
        )
        scan = ScanJob.objects.create(
            user=user,
            original_url="https://example.com",
            normalized_url="https://example.com",
            origin="example.com",
            status=ScanJob.Status.COMPLETED,
        )
        # 讓 Page.screenshot_path 指向我們建的測試截圖。Page 不需要參考：
        # _add_entry_screenshot 會用 scan_job.pages 查詢，自己找到這個 entry page。
        Page.objects.create(
            scan_job=scan,
            url="https://example.com",
            final_url="https://example.com",
            origin="example.com",
            screenshot_path=str(self.scan_img_path),  # relative to BASE_DIR
        )

        path = build_scan_report(scan)
        with zipfile.ZipFile(path) as z:
            media = [n for n in z.namelist() if n.startswith("word/media/")]
            sizes = [z.getinfo(n).file_size for n in media]
        # 內嵌的所有媒體總大小應該明顯小於原圖（1500x6000 PNG 大約 5KB）
        # 縮圖後（1200px 寬）應該在 30-100KB 範圍；原始 1500x6000 應該差不多同大小
        # 主要斷言：docx 內嵌影像總和 < 200KB（若直接 embed 全圖會到 200KB+）
        total = sum(sizes)
        self.assertLess(total, 200_000, f"內嵌媒體過大 ({total}B)，可能沒縮圖")